from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT_FILE = "4.4.4.5.docx"
SCREENSHOT_DIR = Path("screenshots")


customer_44 = [
    (
        "Halaman Registrasi",
        "Halaman registrasi dirancang sebagai pintu masuk bagi calon pelanggan untuk membuat akun pada sistem DIX Game Farm. Antarmuka halaman ini memuat formulir pendaftaran yang terdiri atas nama lengkap, alamat surat elektronik, nomor telepon, password, konfirmasi password, dan alamat lengkap. Susunan komponen dibuat terpusat dalam satu panel formulir agar pengguna fokus pada proses pengisian data. Tombol utama yang disediakan adalah tombol pembuatan akun, sedangkan tautan menuju halaman login ditempatkan sebagai navigasi alternatif bagi pengguna yang telah memiliki akun. Alur penggunaan halaman dimulai dari pengisian seluruh identitas dasar pelanggan, dilanjutkan dengan validasi data oleh sistem, kemudian data dikirim untuk membentuk akun baru.",
    ),
    (
        "Halaman Login Pelanggan",
        "Halaman login pelanggan dirancang untuk memfasilitasi proses autentikasi sebelum pelanggan mengakses fitur transaksi. Komponen utamanya meliputi kolom email, kolom password, tombol masuk, fitur tampil atau sembunyikan password, serta tautan menuju registrasi. Tata letak dibuat sederhana dan terpusat agar proses login dapat dilakukan secara cepat dan jelas. Tombol Masuk berfungsi sebagai aksi utama untuk memeriksa kredensial pengguna. Alur penggunaannya dimulai ketika pelanggan memasukkan email dan password, lalu sistem melakukan verifikasi dan mengarahkan pengguna ke area akun pelanggan apabila autentikasi berhasil.",
    ),
    (
        "Halaman Beranda",
        "Halaman beranda dirancang sebagai antarmuka awal yang menampilkan identitas usaha DIX Game Farm dan mengarahkan pengunjung ke fitur utama sistem. Antarmuka ini memuat bagian hero, deskripsi singkat usaha, tombol menuju katalog dan kontak, penjelasan keunggulan layanan, ajakan pendaftaran, serta informasi ringkas mengenai aktivitas usaha. Tata letak disusun secara bertahap dari pengenalan usaha menuju ajakan melakukan interaksi dengan sistem. Tombol Lihat Katalog dan Hubungi Kami menjadi komponen navigasi utama yang memandu pengunjung menuju proses penelusuran produk atau komunikasi dengan pihak farm. Alur penggunaan dimulai dari pembacaan informasi umum, kemudian pengguna dapat memilih melanjutkan ke katalog atau melakukan pendaftaran akun.",
    ),
    (
        "Halaman Katalog Produk",
        "Halaman katalog produk dirancang sebagai media utama bagi pelanggan untuk menelusuri produk yang ditawarkan. Halaman ini menampilkan filter pencarian berdasarkan kata kunci dan kategori, kemudian dilanjutkan dengan deretan kartu produk yang memuat foto, nama produk, kategori, harga, informasi usia atau berat, serta status ketersediaan stok. Tata letak menggunakan pola grid agar banyak produk dapat dilihat secara sistematis. Tombol Lihat Detail dan tombol penambahan ke keranjang berfungsi sebagai aksi utama untuk melanjutkan proses pembelian. Alur penggunaan dimulai dari penyaringan produk sesuai kebutuhan, dilanjutkan dengan membaca informasi setiap kartu, lalu pelanggan memilih membuka detail produk atau menambahkan produk ke keranjang.",
    ),
    (
        "Halaman Detail Produk",
        "Halaman detail produk dirancang untuk menampilkan informasi produk secara lebih lengkap sebelum pembelian dilakukan. Komponen utama yang disajikan meliputi foto produk, nama produk, kategori, harga, informasi usia atau berat, status stok, keterangan pre-order jika tersedia, deskripsi produk, dan pengatur jumlah pembelian. Tata letak dibagi menjadi area visual produk dan area informasi beserta aksi pembelian sehingga pelanggan dapat membaca informasi dengan lebih terarah. Tombol Tambah ke Keranjang menjadi fitur utama pada halaman ini. Alur penggunaannya dimulai dari pelanggan membuka salah satu produk dari katalog, meninjau spesifikasi produk, menentukan jumlah pesanan, kemudian menambahkan produk ke keranjang.",
    ),
    (
        "Halaman Keranjang Belanja",
        "Halaman keranjang belanja dirancang untuk merangkum seluruh produk yang telah dipilih pelanggan sebelum proses checkout. Halaman ini memuat daftar item, foto produk, nama produk, jumlah, subtotal, tombol pengubah kuantitas, tombol penghapusan item, serta panel ringkasan total belanja. Tata letak dibagi dua bagian, yaitu area daftar item pada sisi utama dan panel ringkasan pada sisi kanan. Tombol Lanjut Checkout menjadi komponen utama untuk meneruskan transaksi, sedangkan tombol Lanjut Belanja memungkinkan pelanggan kembali ke katalog. Alur penggunaan dimulai dari peninjauan item yang sudah dipilih, penyesuaian jumlah produk bila diperlukan, lalu pelanggan melanjutkan transaksi ke tahap checkout.",
    ),
    (
        "Halaman Checkout",
        "Halaman checkout dirancang untuk menampung data pemesanan akhir sebelum transaksi dibentuk menjadi pesanan. Halaman ini memuat ringkasan item pesanan, formulir identitas penerima, alamat pengiriman, pilihan tipe pengiriman, catatan tambahan, daftar metode pembayaran manual, serta ringkasan total pembayaran. Tata letak dibuat dua kolom agar data pesanan dan ringkasan biaya dapat dibaca bersamaan. Tombol utama yang digunakan adalah tombol pembuatan pesanan setelah seluruh data wajib dilengkapi. Alur penggunaan dimulai dari pemeriksaan isi pesanan, pengisian data pengiriman, pemilihan metode pembayaran, lalu konfirmasi untuk membentuk pesanan dalam sistem.",
    ),
    (
        "Halaman Dashboard Pelanggan",
        "Halaman dashboard pelanggan dirancang sebagai pusat informasi akun pelanggan setelah proses login berhasil. Antarmuka ini memuat sapaan pengguna, ringkasan statistik pesanan, tombol menuju profil dan katalog, serta tabel daftar pesanan yang menampilkan nomor invoice, tanggal, total, status, dan tombol detail. Tata letak disusun dengan kombinasi komponen ringkasan pada bagian atas dan tabel transaksi pada bagian bawah. Fungsi utama halaman ini adalah membantu pelanggan melihat riwayat pesanan secara cepat dalam satu tempat. Alur penggunaan dimulai ketika pelanggan masuk ke sistem, membaca ringkasan aktivitas transaksi, lalu memilih salah satu pesanan untuk melihat rincian lebih lanjut.",
    ),
    (
        "Halaman Detail Pesanan",
        "Halaman detail pesanan dirancang untuk menampilkan rincian transaksi dan perkembangan status pesanan pelanggan. Komponen yang disajikan meliputi nomor invoice, status pesanan, penanda tahapan proses, daftar item pesanan, riwayat pembayaran, informasi pengiriman, serta panel ringkasan dan aksi lanjutan. Tata letak dibagi ke area rincian utama dan panel ringkasan di sisi samping. Tombol yang ditampilkan menyesuaikan kondisi pesanan, misalnya upload bukti pembayaran, ganti bukti pembayaran, batalkan pesanan, atau tandai selesai. Alur penggunaannya dimulai dari pelanggan memilih invoice pada dashboard, lalu memantau status dan menjalankan tindakan yang sesuai dengan kondisi pesanan.",
    ),
    (
        "Halaman Upload Bukti Pembayaran",
        "Halaman upload bukti pembayaran dirancang untuk memfasilitasi pengiriman bukti transfer oleh pelanggan. Antarmuka halaman ini memuat nomor invoice, total pembayaran, detail rekening tujuan, nominal transfer yang dikunci, komponen unggah file gambar, serta tombol pengiriman bukti pembayaran. Tata letak dibuat fokus pada proses pengunggahan agar pelanggan tidak mengalami kebingungan saat melakukan konfirmasi pembayaran. Tombol unggah menjadi aksi utama, sedangkan tombol kembali digunakan untuk kembali ke rincian pesanan. Alur penggunaan dimulai dari peninjauan nomor invoice dan nominal transfer, pemilihan file bukti pembayaran, lalu pengiriman file ke sistem untuk diverifikasi admin.",
    ),
    (
        "Halaman Profil Akun",
        "Halaman profil akun dirancang untuk mengelola identitas pelanggan dan keamanan akses akun. Halaman ini terdiri atas dua kelompok utama, yaitu formulir pembaruan data pelanggan dan formulir perubahan password. Komponen yang tersedia mencakup nama lengkap, nomor telepon, alamat lengkap, password lama, password baru, dan konfirmasi password baru. Tata letak dua kolom digunakan agar data profil dan pengaturan keamanan dapat dipisahkan secara jelas. Tombol Simpan Perubahan dan Ganti Password menjadi komponen utama pada halaman ini. Alur penggunaannya dimulai dari pelanggan membuka menu profil, memperbarui data yang diperlukan, lalu menyimpan perubahan sesuai kebutuhan.",
    ),
]


admin_44 = [
    (
        "Halaman Login Admin",
        "Halaman login admin pada sistem ini menggunakan antarmuka autentikasi yang sama dengan login pengguna, tetapi proses validasinya akan memeriksa peran admin sebelum mengarahkan pengguna ke panel administrasi. Komponen utamanya terdiri atas input email, input password, tombol masuk, dan fitur tampil atau sembunyikan password. Tata letaknya tetap sederhana dan terpusat agar proses akses ke panel admin berlangsung efisien. Tombol Masuk berfungsi sebagai aksi utama untuk memulai autentikasi admin. Alur penggunaan dimulai ketika admin memasukkan kredensial, lalu sistem memverifikasi akun dan mengarahkan ke dashboard admin.",
    ),
    (
        "Halaman Dashboard Admin",
        "Halaman dashboard admin dirancang sebagai pusat pemantauan aktivitas operasional sistem. Antarmuka halaman menampilkan sidebar navigasi, kartu statistik utama, notifikasi pembayaran yang menunggu verifikasi, tabel pesanan terbaru, dan informasi produk dengan stok rendah. Tata letaknya dibuat padat namun terstruktur agar admin dapat memperoleh gambaran umum kondisi sistem dari satu halaman. Tombol dan tautan navigasi pada sidebar berfungsi sebagai akses cepat menuju modul pengelolaan. Alur penggunaannya dimulai dari admin membuka dashboard, meninjau indikator operasional, kemudian berpindah ke modul yang membutuhkan tindakan lanjutan.",
    ),
    (
        "Halaman Manajemen Produk",
        "Halaman manajemen produk dirancang untuk mendukung pengelolaan data produk yang dijual pada sistem. Komponen utamanya meliputi tabel daftar produk, foto produk, kategori, harga, stok tersedia, stok dibooking, status aktif, penanda pre-order, serta modal formulir tambah dan edit produk. Tata letak difokuskan pada tabel utama agar admin dapat meninjau seluruh data produk secara ringkas. Tombol Tambah Produk, Edit, dan Hapus menjadi fitur utama pada halaman ini, sedangkan pengelolaan stok dilakukan melalui data produk yang tersedia. Alur penggunaannya dimulai dari peninjauan data produk, lalu admin memilih menambah, mengubah, atau menghapus produk sesuai kebutuhan pengelolaan.",
    ),
    (
        "Halaman Manajemen Pesanan",
        "Halaman manajemen pesanan dirancang sebagai daftar transaksi pelanggan yang dapat dipantau admin secara menyeluruh. Komponen utama halaman ini berupa filter status dan tabel pesanan yang menampilkan invoice, identitas pelanggan, tanggal transaksi, total pesanan, status, tipe pengiriman, dan tombol aksi. Tata letak berbasis tabel digunakan agar proses pemantauan dan seleksi data pesanan menjadi efisien. Fitur utama yang tersedia adalah penyaringan status, akses ke detail pesanan, dan pembatalan pesanan apabila masih diizinkan oleh alur sistem. Alur penggunaan dimulai dari pemilihan filter, dilanjutkan dengan peninjauan daftar pesanan, lalu admin membuka salah satu pesanan untuk pengelolaan lebih lanjut.",
    ),
    (
        "Halaman Detail Pesanan Admin",
        "Halaman detail pesanan admin dirancang untuk menampilkan rincian transaksi secara lengkap sekaligus menjadi pusat tindakan operasional terhadap pesanan tertentu. Komponen yang disajikan meliputi informasi invoice, data akun pelanggan, detail pengiriman saat checkout, tabel item pesanan, riwayat pembayaran, bukti pembayaran, dan panel aksi admin. Tata letaknya memisahkan rincian transaksi di sisi utama dan panel tindakan di sisi kanan agar proses pembacaan data dan pengambilan keputusan dapat dilakukan bersamaan. Fitur utama pada halaman ini adalah perubahan status pesanan, pembatalan pesanan, verifikasi pembayaran, serta pengisian kode resi. Alur penggunaannya dimulai dari pembukaan data pesanan tertentu, lalu admin memeriksa detail transaksi dan menjalankan aksi yang diperlukan.",
    ),
    (
        "Halaman Verifikasi Pembayaran",
        "Halaman verifikasi pembayaran dirancang untuk memusatkan seluruh bukti pembayaran pelanggan yang masih menunggu pemeriksaan. Komponen utamanya berupa tabel pembayaran yang menampilkan nama pelanggan, invoice, bank, nominal transfer, total pesanan, waktu unggah, tombol lihat bukti, serta tombol setujui dan tolak. Tata letak berbasis tabel memudahkan admin membandingkan data pembayaran secara cepat. Fitur utama halaman ini adalah proses verifikasi pembayaran secara manual oleh admin. Alur penggunaannya dimulai ketika admin meninjau bukti transfer yang masuk, memeriksa kecocokan nominal dan bukti, lalu menetapkan keputusan verifikasi.",
    ),
    (
        "Halaman Laporan Penjualan",
        "Halaman laporan penjualan dirancang untuk menyajikan data penjualan berdasarkan rentang waktu tertentu. Antarmuka ini memuat filter tanggal, tombol ekspor data, kartu ringkasan total pesanan dan total pendapatan, tabel daftar pesanan, serta tabel produk terlaris. Tata letaknya disusun dari pengaturan periode, ringkasan statistik, lalu detail data penjualan agar proses analisis berjalan sistematis. Tombol Filter dan Export CSV menjadi komponen utama yang mendukung pengolahan laporan. Alur penggunaan dimulai dari admin menentukan rentang tanggal, kemudian sistem menampilkan data penjualan sesuai periode yang dipilih.",
    ),
    (
        "Halaman Rekap Stok",
        "Halaman rekap stok dirancang untuk memberikan gambaran kondisi persediaan produk secara menyeluruh. Halaman ini menampilkan kartu ringkasan nilai stok, jumlah stok rendah, jumlah stok habis, serta tabel detail stok produk yang memuat kategori, stok tersedia, stok dibooking, harga satuan, nilai stok, dan status persediaan. Tata letak yang digunakan menempatkan indikator ringkasan di bagian atas dan tabel analisis pada bagian bawah. Fitur utama halaman ini adalah pemantauan stok untuk mendukung pengambilan keputusan restok. Alur penggunaan dimulai dari pembacaan indikator ringkas, kemudian admin menelaah detail produk yang perlu mendapat perhatian.",
    ),
    (
        "Halaman Data Pelanggan",
        "Halaman data pelanggan dirancang untuk menampilkan informasi pelanggan beserta ringkasan aktivitas transaksinya. Komponen utamanya meliputi kartu ringkasan jumlah pelanggan, pelanggan aktif, total pesanan, total pendapatan pelanggan, kolom pencarian, dan tabel data pelanggan. Tata letak mengutamakan ringkasan statistik pada bagian atas dan tabel pencarian pada bagian bawah agar proses analisis data pelanggan berjalan efisien. Fitur pencarian membantu admin menemukan pelanggan tertentu secara cepat. Alur penggunaan dimulai dari pembacaan statistik umum, lalu admin dapat menggunakan kolom pencarian atau menelaah tabel pelanggan sesuai kebutuhan.",
    ),
]


customer_45 = [
    (
        "Halaman Registrasi",
        "Halaman registrasi telah diimplementasikan pada sistem untuk mendukung pembuatan akun pelanggan baru. Pada halaman ini ditampilkan formulir pendaftaran yang berisi nama lengkap, email, nomor telepon, password, konfirmasi password, dan alamat lengkap. Pengguna dapat mengirimkan formulir melalui tombol pembuatan akun, kemudian sistem melakukan validasi terhadap setiap input sebelum menyimpan data pelanggan ke basis data. Setelah registrasi berhasil, akun pelanggan dapat langsung digunakan untuk mengakses fitur pemesanan. Screenshot halaman ini menampilkan panel formulir pendaftaran sebagai antarmuka awal pembentukan akun pelanggan.",
    ),
    (
        "Halaman Login Pelanggan",
        "Halaman login pelanggan telah diimplementasikan untuk proses autentikasi pengguna sebelum memasuki fitur transaksi. Halaman ini menampilkan kolom email dan password, tombol Masuk, fitur untuk menampilkan password, serta tautan menuju registrasi. Sistem akan memverifikasi kredensial yang dimasukkan dan mengarahkan pengguna ke dashboard pelanggan apabila data sesuai. Halaman ini berperan sebagai penghubung antara antarmuka publik dan area transaksi pelanggan. Screenshot yang disisipkan menunjukkan form login yang digunakan pelanggan untuk memulai sesi penggunaan sistem.",
    ),
    (
        "Halaman Beranda",
        "Halaman beranda telah diimplementasikan sebagai halaman muka sistem yang menampilkan identitas DIX Game Farm dan navigasi menuju fitur utama. Halaman ini memuat informasi singkat usaha, tombol menuju katalog, tombol kontak, dan elemen promosi layanan. Fungsinya adalah memperkenalkan sistem kepada pengunjung sekaligus mengarahkan mereka menuju proses penelusuran produk atau pendaftaran akun. Dari halaman ini pengguna dapat mulai berinteraksi dengan sistem melalui navigasi utama yang disediakan. Screenshot halaman beranda memperlihatkan tampilan awal aplikasi berbasis website yang digunakan oleh DIX Game Farm.",
    ),
    (
        "Halaman Katalog Produk",
        "Halaman katalog produk telah diimplementasikan untuk menampilkan daftar produk yang tersedia pada sistem. Di dalamnya terdapat filter pencarian, pilihan kategori, daftar kartu produk, informasi harga, status stok, dan tombol untuk melihat detail produk atau menambahkan produk ke keranjang. Halaman ini menjadi sarana utama pelanggan dalam menelusuri produk yang akan dipesan. Sistem juga menampilkan perbedaan kondisi stok, termasuk produk yang tersedia sebagai pre-order. Screenshot halaman ini menunjukkan daftar produk beserta komponen pencarian dan aksi pembelian yang digunakan pelanggan.",
    ),
    (
        "Halaman Detail Produk",
        "Halaman detail produk telah diimplementasikan untuk memberikan informasi yang lebih rinci tentang produk yang dipilih pelanggan. Halaman ini menampilkan foto produk, nama produk, kategori, harga, informasi stok, deskripsi produk, dan komponen pengatur jumlah pesanan. Pelanggan dapat langsung menggunakan tombol Tambah ke Keranjang untuk memasukkan produk ke transaksi. Halaman ini berfungsi sebagai titik keputusan sebelum pelanggan melanjutkan pembelian. Screenshot halaman detail produk memperlihatkan informasi produk secara lengkap beserta tombol aksi pembelian.",
    ),
    (
        "Halaman Keranjang Belanja",
        "Halaman keranjang belanja telah diimplementasikan untuk menampung produk yang dipilih pelanggan sebelum proses checkout dilakukan. Halaman ini menampilkan daftar item, kuantitas masing-masing produk, subtotal, tombol perubahan jumlah, tombol penghapusan item, dan ringkasan total belanja. Pelanggan dapat meninjau kembali produk yang dipilih dan melakukan penyesuaian sebelum membentuk pesanan akhir. Tombol Lanjut Checkout menjadi aksi utama untuk mengarahkan pelanggan menuju tahap pemesanan. Screenshot halaman ini memperlihatkan susunan item dalam keranjang dan panel ringkasan biaya belanja.",
    ),
    (
        "Halaman Checkout",
        "Halaman checkout telah diimplementasikan untuk memfasilitasi pembentukan pesanan akhir oleh pelanggan. Halaman ini menampilkan ringkasan item yang akan dibeli, formulir identitas penerima, data kontak, pilihan tipe pengiriman, catatan tambahan, metode pembayaran manual, dan ringkasan total pembayaran. Sistem hanya mengaktifkan proses pembuatan pesanan setelah data wajib dilengkapi oleh pelanggan. Halaman ini berfungsi sebagai pusat finalisasi transaksi sebelum pesanan direkam ke dalam sistem. Screenshot halaman checkout memperlihatkan formulir pemesanan, daftar metode pembayaran, serta panel ringkasan pembayaran dalam satu tampilan.",
    ),
    (
        "Halaman Dashboard Pelanggan",
        "Halaman dashboard pelanggan telah diimplementasikan sebagai pusat informasi transaksi bagi pengguna yang telah login. Halaman ini menampilkan sapaan pelanggan, statistik jumlah pesanan, serta tabel daftar pesanan yang berisi nomor invoice, tanggal, total, status, dan tombol detail. Fungsi utama halaman ini adalah memberikan ringkasan riwayat transaksi pelanggan dalam satu tampilan. Dari halaman ini pelanggan dapat memilih salah satu pesanan untuk melihat rincian status dan pembayaran secara lebih mendalam. Screenshot halaman dashboard pelanggan menunjukkan statistik transaksi dan daftar pesanan yang sudah tercatat pada sistem.",
    ),
    (
        "Halaman Detail Pesanan",
        "Halaman detail pesanan telah diimplementasikan untuk menampilkan status transaksi pelanggan secara rinci. Komponen yang ditampilkan mencakup invoice, status pesanan, penanda alur status, rincian item pesanan, riwayat pembayaran, informasi pengiriman, dan panel ringkasan. Halaman ini juga menjadi tempat pelanggan menjalankan aksi lanjutan sesuai kondisi pesanan, seperti melihat pembayaran yang telah diverifikasi atau menindaklanjuti proses pesanan. Fungsi halaman ini sangat penting karena menghubungkan riwayat pesanan dengan perkembangan status aktual di sistem. Screenshot halaman detail pesanan memperlihatkan rincian transaksi, status pesanan, dan ringkasan pembayaran pelanggan.",
    ),
    (
        "Halaman Upload Bukti Pembayaran",
        "Halaman upload bukti pembayaran telah diimplementasikan untuk menerima file bukti transfer dari pelanggan terhadap pesanan yang sudah dibuat. Halaman ini menampilkan nomor invoice, total pembayaran, detail rekening tujuan, nominal transfer yang dikunci, komponen unggah gambar, dan tombol kirim bukti pembayaran. Sistem memvalidasi jenis dan ukuran file sebelum bukti dikirim untuk diperiksa admin. Halaman ini berperan sebagai penghubung antara transaksi checkout dan verifikasi pembayaran manual. Screenshot halaman ini memperlihatkan formulir unggah bukti pembayaran beserta informasi invoice dan nominal yang harus dibayarkan.",
    ),
    (
        "Halaman Profil Akun",
        "Halaman profil akun telah diimplementasikan untuk mendukung pemeliharaan data identitas pelanggan dan perubahan password. Pada halaman ini pelanggan dapat memperbarui nama lengkap, nomor telepon, alamat, serta mengganti password melalui formulir yang terpisah. Sistem menyediakan validasi untuk memastikan data yang dimasukkan sesuai dengan ketentuan. Fungsi halaman ini adalah menjaga agar data pelanggan yang digunakan dalam transaksi tetap akurat dan mutakhir. Screenshot halaman profil akun menunjukkan form pengelolaan identitas dan keamanan akun pelanggan.",
    ),
]


admin_45 = [
    (
        "Halaman Login Admin",
        "Halaman login admin telah diimplementasikan melalui form autentikasi yang sama dengan login pengguna, tetapi sistem akan mengarahkan akun dengan peran admin ke panel administrasi. Form ini berisi email, password, tombol masuk, dan fitur tampil atau sembunyikan password. Fungsi utamanya adalah memastikan hanya pengguna dengan hak akses admin yang dapat memasuki area pengelolaan sistem. Setelah kredensial valid, admin diarahkan ke dashboard administrasi. Screenshot halaman ini memperlihatkan antarmuka login yang digunakan admin untuk masuk ke sistem.",
    ),
    (
        "Halaman Dashboard Admin",
        "Halaman dashboard admin telah diimplementasikan sebagai pusat pemantauan aktivitas operasional sistem. Halaman ini menampilkan kartu statistik pesanan hari ini, pembayaran yang menunggu verifikasi, stok rendah, dan total pendapatan, serta dilengkapi tabel pesanan terbaru. Dashboard ini membantu admin mengetahui kondisi sistem secara cepat sebelum masuk ke modul pengelolaan tertentu. Fungsi utamanya adalah memberikan gambaran menyeluruh mengenai transaksi dan stok yang berjalan. Screenshot halaman dashboard admin menunjukkan ringkasan statistik dan daftar pesanan terbaru pada sistem.",
    ),
    (
        "Halaman Manajemen Produk",
        "Halaman manajemen produk telah diimplementasikan untuk mengelola data produk yang ada pada sistem. Halaman ini memuat tabel daftar produk, informasi harga, kategori, stok tersedia, stok dibooking, status produk, serta tombol tambah, edit, dan hapus. Admin juga dapat menambahkan foto produk dan mengatur keterangan pre-order melalui formulir modal. Fungsi halaman ini tidak hanya mengelola data barang, tetapi juga mendukung pengendalian stok melalui perubahan data produk yang tersedia. Screenshot halaman ini memperlihatkan tabel produk dan komponen aksi administrasi yang digunakan untuk pemeliharaan data produk.",
    ),
    (
        "Halaman Manajemen Pesanan",
        "Halaman manajemen pesanan telah diimplementasikan sebagai modul untuk memantau transaksi pelanggan secara keseluruhan. Halaman ini menampilkan filter status dan tabel pesanan yang berisi invoice, pelanggan, tanggal, total, status pesanan, tipe pengiriman, dan tombol detail. Admin menggunakan halaman ini untuk menyeleksi transaksi yang perlu ditindaklanjuti berdasarkan statusnya. Fungsi utama halaman ini adalah memudahkan pengawasan daftar pesanan dalam jumlah banyak. Screenshot halaman ini memperlihatkan tabel manajemen pesanan yang digunakan admin untuk meninjau transaksi pelanggan.",
    ),
    (
        "Halaman Detail Pesanan Admin",
        "Halaman detail pesanan admin telah diimplementasikan untuk menampilkan rincian transaksi sekaligus menjadi pusat pengendalian status pesanan. Halaman ini menampilkan informasi invoice, data pelanggan, detail pengiriman, item pesanan, riwayat pembayaran, dan panel aksi admin untuk memperbarui status atau membatalkan pesanan. Admin juga dapat melakukan tindakan lanjutan seperti memverifikasi pembayaran atau mengisi kode resi sesuai kondisi transaksi. Halaman ini sangat penting karena semua pengelolaan pesanan secara rinci dipusatkan pada satu tampilan. Screenshot halaman ini memperlihatkan detail pesanan beserta panel aksi admin yang digunakan untuk pengendalian status transaksi.",
    ),
    (
        "Halaman Verifikasi Pembayaran",
        "Halaman verifikasi pembayaran telah diimplementasikan sebagai daftar bukti pembayaran yang menunggu pemeriksaan admin. Halaman ini menampilkan nama pelanggan, invoice, bank, nominal transfer, total pesanan, waktu unggah, serta tombol untuk melihat bukti dan menetapkan keputusan verifikasi. Admin dapat menyetujui atau menolak pembayaran berdasarkan bukti transfer yang diunggah pelanggan. Fungsi halaman ini adalah memastikan pembayaran manual tercatat dan tervalidasi sebelum pesanan dilanjutkan ke tahap berikutnya. Screenshot halaman ini menunjukkan tabel pembayaran pending yang digunakan admin dalam proses verifikasi manual.",
    ),
    (
        "Halaman Laporan Penjualan",
        "Halaman laporan penjualan telah diimplementasikan untuk menyajikan rekap penjualan berdasarkan periode tertentu. Halaman ini memuat filter tanggal, kartu total pesanan dan total pendapatan, tabel daftar pesanan, tabel produk terlaris, serta fasilitas ekspor CSV. Admin dapat memanfaatkan halaman ini untuk melakukan analisis transaksi yang telah selesai. Fungsi utamanya adalah menghasilkan laporan penjualan yang terstruktur dan mudah dibaca. Screenshot halaman ini memperlihatkan antarmuka laporan penjualan dengan filter periode dan tabel rekap penjualan.",
    ),
    (
        "Halaman Rekap Stok",
        "Halaman rekap stok telah diimplementasikan untuk menampilkan kondisi persediaan produk secara menyeluruh. Halaman ini menampilkan nilai total stok, jumlah stok rendah, jumlah stok habis, serta tabel detail stok setiap produk. Informasi tersebut membantu admin memonitor persediaan dan menentukan prioritas pengelolaan stok. Fungsi utama halaman ini adalah mendukung pengawasan persediaan secara kuantitatif. Screenshot halaman ini menunjukkan ringkasan indikator stok dan tabel detail persediaan produk yang tersedia pada sistem.",
    ),
    (
        "Halaman Data Pelanggan",
        "Halaman data pelanggan telah diimplementasikan untuk menyajikan daftar pelanggan beserta statistik aktivitas transaksinya. Halaman ini memuat kartu ringkasan jumlah pelanggan, pelanggan aktif, total pesanan, total pendapatan pelanggan, kolom pencarian, dan tabel data pelanggan. Admin dapat menggunakan halaman ini untuk mengenali pelanggan yang aktif serta melihat nilai transaksi kumulatif yang telah dilakukan. Fungsi halaman ini adalah membantu pengelolaan informasi pelanggan pada sisi administrasi. Screenshot halaman ini memperlihatkan tabel data pelanggan dan statistik ringkas pelanggan pada sistem.",
    ),
]


SCREENSHOT_MAP = {
    "Halaman Registrasi": "customer_register.png",
    "Halaman Login Pelanggan": "customer_login.png",
    "Halaman Beranda": "customer_home.png",
    "Halaman Katalog Produk": "customer_catalog.png",
    "Halaman Detail Produk": "customer_product_detail.png",
    "Halaman Keranjang Belanja": "customer_cart.png",
    "Halaman Checkout": "customer_checkout.png",
    "Halaman Dashboard Pelanggan": "customer_dashboard.png",
    "Halaman Detail Pesanan": "customer_order_detail.png",
    "Halaman Upload Bukti Pembayaran": "customer_upload_payment.png",
    "Halaman Profil Akun": "customer_profile.png",
    "Halaman Login Admin": "admin_login.png",
    "Halaman Dashboard Admin": "admin_dashboard.png",
    "Halaman Manajemen Produk": "admin_product_management.png",
    "Halaman Manajemen Pesanan": "admin_order_management.png",
    "Halaman Detail Pesanan Admin": "admin_order_detail.png",
    "Halaman Verifikasi Pembayaran": "admin_payment_verification.png",
    "Halaman Laporan Penjualan": "admin_sales_report.png",
    "Halaman Rekap Stok": "admin_stock_report.png",
    "Halaman Data Pelanggan": "admin_customer_report.png",
}


figure_counter = 1


def set_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def configure_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)


def add_center_heading(doc: Document, text: str, size: int = 14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=size, bold=True)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)


def add_left_heading(doc: Document, text: str, size: int = 12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size=size, bold=True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def add_body(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text.strip())
    set_font(run, size=12)


def add_item(doc: Document, letter: str, title: str, body: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run1 = p.add_run(f"{letter}. ")
    set_font(run1, size=12, bold=True)
    run2 = p.add_run(title)
    set_font(run2, size=12, bold=True)
    add_body(doc, body)


def add_section_intro(doc: Document, text: str):
    add_body(doc, text)


def add_image(doc: Document, title: str):
    global figure_counter
    filename = SCREENSHOT_MAP.get(title)
    if not filename:
        return

    image_path = SCREENSHOT_DIR / filename
    if not image_path.exists():
        return

    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(str(image_path), width=Inches(6.2), height=Inches(3.49))
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(2)

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = p_cap.add_run(f"Gambar 4.{figure_counter} {title}")
    set_font(cap, size=11, italic=True)
    p_cap.paragraph_format.space_after = Pt(8)
    figure_counter += 1


def build_document():
    doc = Document()
    configure_document(doc)

    add_center_heading(doc, "BAB IV")
    add_center_heading(doc, "HASIL DAN PEMBAHASAN")

    add_left_heading(doc, "4.4 Perancangan Antarmuka")
    add_section_intro(
        doc,
        "Bagian ini menjelaskan perancangan antarmuka sistem berdasarkan halaman yang benar-benar tersedia pada aplikasi DIX Game Farm. Uraian difokuskan pada tujuan halaman, susunan komponen utama, tata letak, fitur yang disediakan, dan alur penggunaan oleh masing-masing peran dalam sistem.",
    )

    add_left_heading(doc, "4.4.1 Halaman Pelanggan")
    for idx, (title, body) in enumerate(customer_44):
        add_item(doc, chr(ord("a") + idx), title, body)
        add_image(doc, title)

    add_left_heading(doc, "4.4.2 Halaman Admin")
    for idx, (title, body) in enumerate(admin_44):
        add_item(doc, chr(ord("a") + idx), title, body)
        add_image(doc, title)

    doc.add_section(WD_SECTION_START.NEW_PAGE)

    add_left_heading(doc, "4.5 Implementasi Sistem")
    add_section_intro(
        doc,
        "Bagian ini memaparkan hasil realisasi antarmuka dan fungsi halaman pada sistem yang telah berjalan. Penjelasan difokuskan pada komponen yang tampil, aksi yang dapat dilakukan pengguna, alur penggunaan, serta dokumentasi visual dari halaman yang diimplementasikan.",
    )

    add_left_heading(doc, "4.5.1 Halaman Pelanggan")
    for idx, (title, body) in enumerate(customer_45):
        add_item(doc, chr(ord("a") + idx), title, body)
        add_image(doc, title)

    add_left_heading(doc, "4.5.2 Halaman Admin")
    for idx, (title, body) in enumerate(admin_45):
        add_item(doc, chr(ord("a") + idx), title, body)
        add_image(doc, title)

    doc.save(OUTPUT_FILE)
    print(f"Generated {OUTPUT_FILE} with corrected screenshots")


if __name__ == "__main__":
    build_document()
