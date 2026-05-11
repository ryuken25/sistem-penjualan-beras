from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT_FILE = "4.4.4.5.docx"


def set_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def configure_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)


def add_center_heading(doc: Document, text: str, size: int = 14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=size, bold=True)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)


def add_left_heading(doc: Document, text: str, size: int = 12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size=size, bold=True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def add_body(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text.strip())
    set_font(run, size=12, bold=False)


def add_item(doc: Document, letter: str, title: str, body: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run1 = p.add_run(f"{letter}. ")
    set_font(run1, size=12, bold=True)
    run2 = p.add_run(title)
    set_font(run2, size=12, bold=True)
    add_body(doc, body)


customer_44 = [
    (
        "Halaman Registrasi",
        "Halaman Registrasi dirancang sebagai antarmuka awal bagi calon pelanggan untuk membuat akun pada sistem DIX Game Farm. Pada halaman ini disediakan formulir yang memuat nama lengkap, email, nomor HP, password, konfirmasi password, dan alamat lengkap sehingga data dasar pelanggan dapat direkam sejak awal. Tata letak halaman dibuat terpusat dengan panel formulir yang dominan, identitas logo sistem, serta penanda bidang wajib agar mudah dipahami oleh pengguna baru. Tombol utama berupa Buat Akun ditempatkan pada bagian bawah formulir, sedangkan tautan menuju halaman login disediakan sebagai navigasi alternatif bagi pengguna yang sudah memiliki akun. Alur penggunaan halaman dimulai dari pengisian seluruh data, dilanjutkan validasi isian, lalu pengiriman formulir untuk pembuatan akun pelanggan.",
    ),
    (
        "Halaman Login",
        "Halaman Login dirancang untuk memfasilitasi proses autentikasi pelanggan sebelum mengakses fitur transaksi. Elemen antarmuka utamanya terdiri atas identitas sistem, kolom email, kolom password, fitur tampil atau sembunyikan password, tombol masuk, dan tautan ke halaman pendaftaran. Layout halaman menggunakan panel tunggal yang dipusatkan agar fokus pengguna langsung tertuju pada formulir autentikasi. Tombol Masuk berperan sebagai aksi utama, sedangkan tautan Daftar Sekarang mendukung pengguna baru yang belum memiliki akun. Alur penggunaan halaman dimulai ketika pelanggan memasukkan email dan password, kemudian sistem memeriksa kredensial sebelum mengarahkan pengguna ke area akun pelanggan.",
    ),
    (
        "Halaman Beranda",
        "Halaman Beranda dirancang sebagai antarmuka publik yang memperkenalkan layanan pemesanan ayam pada DIX Game Farm sekaligus mengarahkan pengunjung menuju fitur utama sistem. Komponen yang ditampilkan meliputi bagian hero berisi judul, deskripsi singkat, logo, tombol menuju katalog dan kontak, bagian keunggulan layanan, ajakan pendaftaran, serta ringkasan statistik usaha. Tata letak disusun secara vertikal dari pengenalan, alasan memilih layanan, hingga ajakan melakukan transaksi sehingga alur pembacaan pengguna menjadi runtut. Tombol Lihat Katalog dan Hubungi Kami ditempatkan pada area awal sebagai titik masuk ke proses bisnis. Alur penggunaan dimulai dari eksplorasi informasi umum, kemudian pengguna dapat melanjutkan ke katalog produk atau melakukan pendaftaran akun.",
    ),
    (
        "Halaman Katalog Produk",
        "Halaman Katalog Produk dirancang sebagai antarmuka utama pelanggan untuk menelusuri produk yang tersedia pada DIX Game Farm. Bagian atas halaman memuat formulir pencarian berdasarkan kata kunci dan kategori, sedangkan bagian utama menampilkan kartu produk berisi foto, nama produk, kategori, informasi usia atau berat, harga, status ketersediaan, serta tombol lihat detail atau tambah ke keranjang. Tata letak menggunakan model grid agar banyak produk dapat dilihat secara ringkas namun tetap terstruktur. Tombol Filter digunakan untuk menyaring produk, sedangkan tombol aksi pada setiap kartu mempercepat perpindahan ke proses pembelian. Alur penggunaan halaman dimulai dari pencarian atau pemilihan kategori, dilanjutkan dengan membaca informasi stok dan membuka detail produk atau langsung menambahkan produk ke keranjang.",
    ),
    (
        "Halaman Detail Produk",
        "Halaman Detail Produk dirancang untuk menampilkan informasi produk secara lebih lengkap sebelum pelanggan memutuskan pembelian. Halaman ini memuat foto produk berukuran besar, nama produk, kategori, usia atau berat, harga, keterangan ketersediaan stok, penanda pre-order apabila ada, serta deskripsi produk. Tata letak dibagi dua kolom, yaitu area visual produk di sisi kiri dan area informasi serta aksi pembelian di sisi kanan sehingga fokus pengguna tetap terarah. Fitur utama yang disediakan berupa pengatur jumlah pesanan dan tombol tambah ke keranjang, disertai tombol kembali ke katalog. Alur penggunaannya dimulai dari peninjauan detail produk, penentuan jumlah pembelian, kemudian pengguna dapat menambahkan produk ke keranjang untuk melanjutkan transaksi.",
    ),
    (
        "Halaman Keranjang Belanja",
        "Halaman Keranjang Belanja dirancang untuk merangkum produk yang telah dipilih pelanggan sebelum checkout. Elemen antarmuka utama mencakup daftar item, gambar produk, nama, kategori, informasi stok atau label pre-order, pengatur jumlah, subtotal per item, tombol hapus, tombol kosongkan keranjang, serta panel ringkasan belanja. Layout dibagi menjadi dua kolom, yaitu area daftar item di sisi utama dan ringkasan total pada sisi samping agar keputusan transaksi dapat dilakukan dengan cepat. Fitur utama halaman ini adalah penyesuaian jumlah barang dan tombol Lanjut Checkout. Alur penggunaannya dimulai dari peninjauan item, perubahan kuantitas bila diperlukan, lalu pengguna dapat melanjutkan pembelian atau kembali ke katalog.",
    ),
    (
        "Halaman Checkout / Pemesanan",
        "Halaman Checkout atau Pemesanan dirancang untuk membantu pelanggan menyelesaikan transaksi setelah memilih produk. Antarmuka halaman menampilkan review item pesanan, formulir informasi penerima dan pengiriman, pilihan tipe pengiriman, catatan pesanan, daftar metode pembayaran, serta ringkasan total pembayaran. Tata letak dibuat dua kolom, dengan kolom utama berisi data transaksi dan kolom samping menampilkan total pembayaran agar informasi penting selalu terlihat. Tombol utama berupa Buat Pesanan ditempatkan pada panel ringkasan, sedangkan tombol kembali ke keranjang disediakan sebagai navigasi koreksi. Alur penggunaannya dimulai dari pemeriksaan item, pengisian data pengiriman dan metode pembayaran, lalu konfirmasi pemesanan hingga pesanan tercatat ke dalam sistem.",
    ),
    (
        "Halaman Dashboard Pelanggan dan Riwayat Pesanan",
        "Halaman Dashboard Pelanggan dirancang sebagai pusat informasi akun yang menampilkan riwayat pesanan dalam satu tampilan. Elemen yang disajikan meliputi sapaan pengguna, kartu statistik jumlah pesanan menurut kondisi umum, tombol menuju profil dan katalog, serta tabel daftar pesanan berisi nomor invoice, tanggal, total, status, dan tombol detail. Tata letak menggabungkan ringkasan visual di bagian atas dengan tabel data di bagian bawah agar pengguna dapat memantau transaksi secara cepat sekaligus terperinci. Tombol Detail menjadi fitur utama untuk melanjutkan penelusuran status pesanan. Alur penggunaan halaman dimulai setelah login, ketika pelanggan melihat ringkasan transaksi, lalu membuka pesanan tertentu untuk melihat perkembangan lebih rinci.",
    ),
    (
        "Halaman Detail Pesanan dan Status Pesanan",
        "Halaman Detail Pesanan dirancang untuk menampilkan perkembangan transaksi secara rinci sejak pesanan dibuat hingga selesai. Komponen utama pada halaman ini mencakup nomor invoice, status pesanan, tanggal transaksi, penanda alur status, instruksi pembayaran, daftar item pesanan, riwayat pembayaran, ringkasan pengiriman, dan tombol aksi sesuai kondisi pesanan. Layout halaman menggunakan pembagian kolom utama untuk detail transaksi dan kolom samping untuk ringkasan serta tindakan lanjutan. Fitur utama yang tersedia meliputi upload bukti pembayaran, ganti bukti pembayaran, batalkan pesanan, dan konfirmasi pesanan selesai. Alur penggunaannya dimulai dari pemilihan pesanan pada dashboard, dilanjutkan dengan pemantauan status, lalu pengguna menjalankan aksi lanjutan sesuai tahap transaksi.",
    ),
    (
        "Halaman Upload Bukti Pembayaran",
        "Halaman Upload Bukti Pembayaran dirancang sebagai antarmuka khusus untuk mengirimkan bukti transfer atas pesanan yang telah dibuat. Halaman ini menampilkan nomor invoice, total pembayaran, informasi rekening atau akun tujuan, nominal transfer yang dikunci, komponen unggah file, serta pratinjau gambar bukti pembayaran. Tata letaknya dibuat sederhana dan terfokus agar pengguna tidak bingung dalam menyelesaikan kewajiban pembayaran. Tombol utama berupa Upload Bukti Pembayaran atau Ganti Bukti Pembayaran ditempatkan pada bagian akhir formulir, disertai tombol kembali ke detail pesanan. Alur penggunaan dimulai dari melihat data tujuan transfer, memilih file bukti pembayaran, memeriksa pratinjau, lalu mengirimkan file untuk proses verifikasi admin.",
    ),
    (
        "Halaman Profil Akun",
        "Halaman Profil Akun dirancang untuk memungkinkan pelanggan mengelola data identitas dan keamanan akunnya. Antarmuka halaman memuat dua kelompok utama, yaitu formulir pembaruan data akun yang berisi nama lengkap, nomor HP, dan alamat, serta formulir perubahan password yang berisi password lama, password baru, dan konfirmasi password. Tata letak dibuat dalam dua kolom agar informasi profil dan pengaturan keamanan dapat dipisahkan namun tetap berada dalam satu halaman. Tombol Simpan Perubahan dan Ganti Password disediakan sebagai aksi utama sesuai fungsi masing-masing formulir. Alur penggunaan halaman dimulai dari pelanggan membuka menu profil, memperbarui data yang diperlukan, kemudian menyimpan perubahan agar data akun tetap sesuai kondisi terbaru.",
    ),
]


admin_44 = [
    (
        "Halaman Login Admin",
        "Halaman Login Admin pada implementasi ini dirancang menggunakan formulir autentikasi yang sama dengan pengguna, namun proses validasi peran akan membawa admin ke panel administrasi setelah kredensial benar. Antarmuka tetap memuat input email, password, fitur tampil atau sembunyikan password, dan tombol masuk. Layout yang sederhana membuat proses akses cepat, sedangkan pemisahan hak akses dilakukan oleh sistem pada tahap autentikasi. Tombol masuk menjadi aksi utama untuk mengawali pengelolaan sistem. Alur penggunaannya dimulai dari admin memasukkan akun, kemudian sistem mengenali peran admin dan mengarahkan pengguna ke dashboard administrasi.",
    ),
    (
        "Halaman Dashboard Admin",
        "Halaman Dashboard Admin dirancang sebagai pusat pemantauan operasional sistem bagi pengelola DIX Game Farm. Antarmuka halaman menampilkan sidebar navigasi, kartu statistik utama, pemberitahuan pembayaran yang menunggu verifikasi, tabel pesanan terbaru, dan daftar produk dengan stok rendah. Tata letak dibangun dengan struktur panel dan tabel agar data penting dapat dibaca cepat oleh admin tanpa berpindah halaman. Fitur utama yang disediakan berupa akses cepat ke verifikasi pembayaran, manajemen pesanan, dan manajemen produk. Alur penggunaan halaman dimulai dari admin membuka dashboard, meninjau indikator operasional, lalu melanjutkan ke menu administrasi yang membutuhkan tindakan.",
    ),
    (
        "Halaman Manajemen Produk dan Stok",
        "Halaman Manajemen Produk dan Stok dirancang untuk membantu admin mengelola data produk sekaligus memantau ketersediaan stok. Elemen antarmuka utamanya mencakup kartu ringkasan produk, tabel daftar produk dengan foto, kategori, harga, stok tersedia, stok dibooking, status aktif, indikator pre-order, serta modal formulir tambah dan edit produk. Tata letak dibuat dengan tabel utama yang luas dan formulir modal agar proses pengelolaan data dapat dilakukan tanpa berpindah konteks. Tombol Tambah Produk, Edit, dan Hapus menjadi aksi utama yang mendukung pengelolaan data barang. Alur penggunaannya dimulai dari peninjauan daftar produk, kemudian admin dapat menambah, mengubah, atau menghapus data sesuai kebutuhan pengelolaan stok ayam dan produk terkait.",
    ),
    (
        "Halaman Manajemen Pesanan",
        "Halaman Manajemen Pesanan dirancang sebagai daftar kerja admin untuk memantau transaksi pelanggan secara menyeluruh. Komponen utamanya berupa form filter status, tabel pesanan dengan informasi invoice, data pelanggan, tanggal pesanan, total transaksi, status, tipe pengiriman, dan tombol aksi. Tata letak halaman memprioritaskan tabel karena halaman ini berfungsi sebagai pusat kontrol daftar pesanan. Fitur utama yang tersedia adalah penyaringan berdasarkan status, pembukaan detail pesanan, dan pembatalan pesanan apabila masih berada pada tahap yang diizinkan. Alur penggunaan dimulai dari admin memilih filter yang diperlukan, meninjau daftar pesanan, lalu membuka detail transaksi untuk tindakan lanjutan.",
    ),
    (
        "Halaman Detail Pesanan dan Perubahan Status",
        "Halaman Detail Pesanan dan Perubahan Status dirancang untuk menampilkan isi transaksi secara lengkap sekaligus menyediakan kontrol operasional bagi admin. Elemen antarmukanya meliputi informasi pesanan, identitas pelanggan, snapshot data pengiriman saat checkout, tabel item pesanan, riwayat pembayaran, tampilan bukti pembayaran, serta panel aksi admin untuk mengubah status, membatalkan pesanan, menyetujui atau menolak pembayaran, dan mengisi kode resi. Tata letak dibagi menjadi area informasi utama dan panel aksi di sisi samping agar data dan keputusan operasional berada dalam satu halaman. Alur penggunaan dimulai dari pemilihan pesanan, pemeriksaan data, lalu admin menjalankan perubahan status sesuai kondisi transaksi yang sedang berlangsung.",
    ),
    (
        "Halaman Verifikasi Pembayaran",
        "Halaman Verifikasi Pembayaran dirancang untuk memusatkan seluruh bukti transfer yang masih menunggu pemeriksaan admin. Antarmuka utama menampilkan tabel pembayaran dengan data pelanggan, nomor invoice, bank, nominal transfer, total pesanan, waktu unggah, tombol melihat bukti, serta tombol setujui dan tolak. Tata letak menempatkan daftar pembayaran pada area utama sehingga admin dapat memeriksa transaksi satu per satu secara efisien. Fitur tambahan berupa panduan verifikasi membantu menjaga konsistensi pemeriksaan. Alur penggunaan halaman dimulai dari admin memilih salah satu pembayaran, memeriksa kesesuaian bukti dan nominal, kemudian memutuskan untuk menyetujui atau menolak pembayaran tersebut.",
    ),
    (
        "Halaman Laporan Penjualan",
        "Halaman Laporan Penjualan dirancang untuk menyajikan ringkasan hasil transaksi yang telah selesai pada periode tertentu. Elemen utama yang disediakan meliputi filter tanggal mulai dan akhir, tombol ekspor data, kartu ringkasan total pesanan dan total pendapatan, tabel daftar pesanan, serta tabel produk terlaris. Tata letak disusun dari kontrol filter di bagian atas, ringkasan statistik di bagian tengah, dan detail data pada bagian bawah agar proses analisis berjalan sistematis. Tombol Filter dan Export CSV berfungsi sebagai fitur utama untuk penyaringan dan keluaran data. Alur penggunaan dimulai dari penentuan rentang tanggal, kemudian admin meninjau hasil penjualan dan performa produk berdasarkan periode yang dipilih.",
    ),
    (
        "Halaman Rekap Stok",
        "Halaman Rekap Stok dirancang untuk memberikan gambaran persediaan produk secara menyeluruh kepada admin. Antarmuka halaman menampilkan kartu ringkasan nilai stok, jumlah stok rendah, jumlah stok habis, serta tabel detail produk yang berisi stok tersedia, stok dibooking, harga satuan, nilai stok, dan status persediaan. Tata letak yang digunakan menempatkan indikator ringkasan di bagian atas dan tabel analisis di bawahnya agar kondisi stok dapat dipahami secara cepat dan mendalam. Fitur utama halaman ini adalah pemantauan ketersediaan barang dan prioritas restok. Alur penggunaannya dimulai dari pembacaan indikator ringkas, lalu admin menelaah detail produk yang memerlukan perhatian.",
    ),
    (
        "Halaman Data Pelanggan",
        "Halaman Data Pelanggan dirancang untuk menampilkan rekap pelanggan yang sudah terdaftar beserta aktivitas transaksinya. Komponen utama pada halaman ini mencakup kartu ringkasan jumlah pelanggan, pelanggan aktif, total pesanan, total pendapatan pelanggan, kolom pencarian, dan tabel data pelanggan berisi nama, email, nomor HP, total pesanan, total belanja, serta tanggal pendaftaran. Tata letak dibuat dengan ringkasan statistik di bagian atas dan tabel pencarian di bagian bawah agar proses peninjauan data lebih efisien. Fitur pencarian membantu admin menemukan pelanggan tertentu secara cepat. Alur penggunaan dimulai dari peninjauan statistik umum, kemudian admin menggunakan kolom pencarian atau membaca tabel untuk menganalisis perilaku pelanggan.",
    ),
]


customer_45 = [
    (
        "Halaman Registrasi",
        "Halaman Registrasi telah diimplementasikan untuk memfasilitasi pembuatan akun pelanggan baru pada sistem DIX Game Farm. Halaman ini menampilkan form yang berisi nama lengkap, email, nomor HP, password, konfirmasi password, dan alamat lengkap, disertai validasi ketika data belum sesuai. Pengguna dapat mengirimkan data melalui tombol Buat Akun, kemudian sistem menyimpan akun sebagai pelanggan dan melanjutkan proses autentikasi awal setelah registrasi berhasil. Alur penggunaan dimulai dari pengisian seluruh form, validasi sistem, lalu akun langsung dapat digunakan untuk mengakses fitur pemesanan. Apabila screenshot halaman ini dimasukkan ke dokumen, tampilan yang terlihat adalah panel registrasi terpusat dengan seluruh bidang identitas pelanggan dan tombol pembuatan akun.",
    ),
    (
        "Halaman Login",
        "Halaman Login telah diimplementasikan sebagai gerbang autentikasi pengguna sebelum masuk ke fitur transaksi. Komponen yang tampil terdiri atas logo sistem, input email, input password, tombol untuk menampilkan atau menyembunyikan password, tombol Masuk, serta tautan menuju registrasi. Melalui halaman ini pengguna dapat melakukan autentikasi, sedangkan sistem akan memeriksa email dan password lalu mengarahkan pengguna ke dashboard sesuai peran yang dimiliki. Alur penggunaan berlangsung dari pengisian kredensial, validasi login, sampai pengguna berhasil masuk ke area pelanggan. Apabila screenshot halaman ini ditampilkan, gambar akan memperlihatkan form login yang sederhana dengan dua isian utama dan tombol Masuk sebagai aksi dominan.",
    ),
    (
        "Halaman Beranda",
        "Halaman Beranda telah direalisasikan sebagai landing page publik yang menampilkan identitas DIX Game Farm sekaligus pintu masuk menuju proses pemesanan. Halaman ini memperlihatkan bagian hero, tombol Lihat Katalog dan Hubungi Kami, kartu keunggulan layanan, ajakan mendaftar, serta ringkasan statistik usaha. Pengguna dapat langsung memilih untuk menelusuri produk atau melakukan pendaftaran akun dari halaman ini. Alur penggunaannya bersifat informatif, yaitu dimulai dari membaca deskripsi usaha, kemudian memilih tombol aksi yang sesuai dengan kebutuhan. Apabila screenshot halaman ini dimasukkan, tampilan akan menunjukkan area hero dengan identitas peternakan, tombol ajakan, dan bagian informasi pendukung di bawahnya.",
    ),
    (
        "Halaman Katalog Produk",
        "Halaman Katalog Produk telah diimplementasikan untuk menampilkan daftar produk yang dapat dipesan pelanggan. Halaman ini memuat form pencarian berdasarkan kata kunci dan kategori, kemudian menampilkan kartu produk dengan foto, nama, kategori, harga, informasi usia atau berat, serta status ketersediaan seperti tersedia, terbatas, habis, atau pre-order. Pengguna dapat membuka detail produk atau menambahkan barang ke keranjang; jika belum login, sistem mengarahkan pengguna ke halaman login terlebih dahulu. Alur penggunaan dimulai dari proses penyaringan produk, dilanjutkan membaca informasi pada kartu, kemudian pengguna memilih detail atau menambah ke keranjang. Apabila screenshot halaman ini ditampilkan, gambar akan memperlihatkan area filter di bagian atas dan deretan kartu produk di area utama.",
    ),
    (
        "Halaman Detail Produk",
        "Halaman Detail Produk telah diimplementasikan untuk menampilkan informasi produk secara lebih rinci sebelum pembelian dilakukan. Komponen yang tampil meliputi foto produk, nama produk, kategori, usia atau berat, harga, keterangan stok, informasi pre-order jika tersedia, deskripsi produk, pengatur jumlah pembelian, dan tombol tambah ke keranjang. Pada halaman ini pelanggan dapat menentukan jumlah pesanan secara langsung, kemudian menambahkan produk ke keranjang melalui tombol aksi utama. Alur penggunaan dimulai saat pelanggan membuka salah satu produk dari katalog, meninjau detail, menentukan jumlah, lalu melanjutkan ke keranjang. Apabila screenshot halaman ini dimasukkan, gambar akan menunjukkan foto produk di satu sisi dan informasi harga, stok, serta tombol pembelian di sisi lainnya.",
    ),
    (
        "Halaman Keranjang Belanja",
        "Halaman Keranjang Belanja telah direalisasikan untuk menampung seluruh produk yang dipilih pelanggan sebelum checkout. Halaman ini menampilkan daftar item lengkap dengan foto, nama, kategori, jumlah, subtotal, tombol tambah atau kurang kuantitas, tombol hapus item, tombol kosongkan keranjang, serta panel ringkasan total belanja. Aksi yang dapat dilakukan pengguna adalah mengubah jumlah pesanan, menghapus item tertentu, mengosongkan seluruh keranjang, atau melanjutkan ke checkout. Alur penggunaan dimulai dari peninjauan kembali item yang dipilih, penyesuaian kuantitas jika diperlukan, lalu pengguna melanjutkan transaksi melalui tombol Lanjut Checkout. Apabila screenshot halaman ini ditampilkan, gambar akan memperlihatkan daftar item belanja di sisi kiri dan ringkasan total transaksi di sisi kanan.",
    ),
    (
        "Halaman Checkout / Pemesanan",
        "Halaman Checkout telah diimplementasikan untuk memfasilitasi pembentukan pesanan akhir oleh pelanggan. Halaman ini menampilkan review item pesanan, formulir data penerima, email, nomor HP, alamat lengkap, pilihan tipe pengiriman, catatan tambahan, pilihan metode pembayaran manual seperti bank dan dompet digital, serta ringkasan total bayar. Pengguna dapat melengkapi seluruh data dan menekan tombol Buat Pesanan setelah semua bidang wajib terisi; sistem juga menahan pengiriman ganda melalui token submit agar pemesanan tidak tercatat dua kali. Alur penggunaan dimulai dari pemeriksaan item, pengisian data pengiriman dan pembayaran, lalu konfirmasi pemesanan sampai nomor invoice terbentuk. Apabila screenshot halaman ini dimasukkan, tampilan akan memperlihatkan form checkout yang lengkap dengan kartu pilihan pembayaran dan panel ringkasan transaksi.",
    ),
    (
        "Halaman Dashboard Pelanggan dan Riwayat Pesanan",
        "Halaman Dashboard Pelanggan telah diimplementasikan sebagai pusat pemantauan riwayat transaksi pelanggan. Halaman ini menampilkan sapaan pengguna, statistik jumlah pesanan menurut kondisi umum, tombol menuju profil dan katalog, serta tabel riwayat pesanan yang memuat invoice, tanggal, total, status, dan tombol detail. Pengguna dapat memanfaatkan halaman ini untuk melihat histori transaksi sekaligus mengetahui apakah pesanan masih menunggu pembayaran, sedang diproses, sedang dikirim, atau telah selesai. Alur penggunaan dimulai segera setelah login ketika sistem menampilkan ringkasan pesanan pengguna, lalu pelanggan memilih salah satu invoice untuk membuka rincian transaksi. Apabila screenshot halaman ini ditampilkan, gambar akan menunjukkan kartu statistik di bagian atas dan tabel riwayat pesanan pada bagian bawah halaman.",
    ),
    (
        "Halaman Detail Pesanan dan Status Pesanan",
        "Halaman Detail Pesanan telah diimplementasikan untuk menampilkan status transaksi secara rinci dan dinamis. Komponen yang tampil mencakup nomor invoice, badge status, penanda alur status sesuai tipe pengiriman, instruksi pembayaran, daftar item, riwayat pembayaran, ringkasan pengiriman, serta tombol aksi seperti upload bukti pembayaran, ganti bukti pembayaran, batalkan pesanan, atau tandai pesanan selesai. Fungsi halaman ini sangat penting karena status pesanan pelanggan dan riwayat pembayaran dipusatkan di sini, termasuk kondisi ketika bukti pembayaran ditolak dan pelanggan diminta mengunggah ulang sebelum batas waktu berakhir. Alur penggunaan dimulai dari dashboard, lalu pelanggan membuka rincian invoice tertentu dan menjalankan aksi yang sesuai dengan tahap transaksi. Apabila screenshot halaman ini dimasukkan, gambar akan memperlihatkan status tracker, rincian item, dan panel ringkasan aksi pelanggan dalam satu tampilan.",
    ),
    (
        "Halaman Upload Bukti Pembayaran",
        "Halaman Upload Bukti Pembayaran telah diimplementasikan untuk menerima bukti transfer pelanggan terhadap pesanan yang sudah terbentuk. Halaman ini menampilkan invoice, total pembayaran, rekening atau akun tujuan, nominal yang dikunci sesuai grand total, komponen unggah file gambar, pratinjau file, serta tombol kirim bukti. Pengguna dapat menggunakan halaman yang sama untuk unggah awal maupun penggantian bukti pembayaran yang masih perlu diperbaiki. Alur penggunaan dimulai dari pelanggan membuka menu upload dari detail pesanan, memilih file bukti transfer, melihat pratinjau, lalu mengirimkan file untuk diverifikasi admin. Apabila screenshot halaman ini ditampilkan, gambar akan memperlihatkan informasi invoice, nominal transfer, detail tujuan pembayaran, dan area unggah gambar bukti transfer.",
    ),
    (
        "Halaman Profil Akun",
        "Halaman Profil Akun telah diimplementasikan untuk mendukung pemeliharaan data identitas pelanggan dan keamanan akses. Halaman ini menampilkan form informasi akun yang dapat diubah, yaitu nama lengkap, nomor HP, dan alamat, serta form terpisah untuk mengganti password dengan verifikasi password lama. Pengguna dapat memperbarui biodata melalui tombol Simpan Perubahan dan mengganti kata sandi melalui tombol Ganti Password. Alur penggunaan dimulai dari pelanggan membuka menu profil, memperbarui data yang sudah tidak sesuai, lalu menyimpan perubahan agar data yang digunakan saat transaksi tetap mutakhir. Apabila screenshot halaman ini dimasukkan, tampilan akan menunjukkan dua panel utama, yaitu panel informasi akun dan panel perubahan password.",
    ),
]


admin_45 = [
    (
        "Halaman Login Admin",
        "Halaman Login Admin telah diimplementasikan melalui form autentikasi yang sama dengan pengguna umum, kemudian sistem melakukan identifikasi peran untuk mengarahkan akun admin ke dashboard administrasi. Komponen yang tampil berupa input email, input password, tombol tampil atau sembunyikan password, dan tombol Masuk. Aksi yang dapat dilakukan pada halaman ini adalah autentikasi admin ke sistem, sedangkan pemisahan hak akses diproses setelah data login tervalidasi. Alur penggunaan dimulai dari pengisian kredensial admin hingga sistem menampilkan dashboard admin sebagai tujuan akhir. Apabila screenshot halaman ini dimasukkan, gambar akan memperlihatkan form login tunggal yang pada implementasinya melayani autentikasi pelanggan maupun admin.",
    ),
    (
        "Halaman Dashboard Admin",
        "Halaman Dashboard Admin telah direalisasikan sebagai halaman pemantauan utama bagi pengelola sistem. Halaman ini menampilkan kartu statistik pesanan hari ini, jumlah pembayaran yang belum diverifikasi, jumlah produk dengan stok rendah, dan total pendapatan; selain itu tersedia ringkasan pesanan terbaru, notifikasi pembayaran menunggu verifikasi, serta tabel produk dengan stok rendah. Admin dapat menggunakan halaman ini untuk menentukan prioritas tindakan, misalnya membuka verifikasi pembayaran atau memeriksa stok menipis. Alur penggunaan dimulai dari dashboard setelah login, kemudian admin melanjutkan ke menu yang membutuhkan penanganan. Apabila screenshot halaman ini ditampilkan, gambar akan menunjukkan kombinasi kartu statistik, alert operasional, dan tabel ringkasan transaksi maupun stok.",
    ),
    (
        "Halaman Manajemen Produk dan Stok",
        "Halaman Manajemen Produk dan Stok telah diimplementasikan pada area admin untuk mengelola produk secara lengkap. Komponen utamanya berupa kartu ringkasan jumlah produk, tabel daftar produk dengan foto, kategori, harga, stok tersedia, stok dibooking, status aktif, dan penanda pre-order, serta modal form untuk tambah dan edit produk yang memuat kategori, nama, usia atau berat, harga, stok, foto, deskripsi, status aktif, dan estimasi pre-order. Admin dapat menambah produk baru, memperbarui data produk, mengelola stok melalui form produk, dan menghapus produk selama tidak ada stok yang sedang dibooking. Alur penggunaan dimulai dari peninjauan tabel produk, lalu admin memilih aksi tambah, ubah, atau hapus sesuai kebutuhan pengelolaan. Apabila screenshot halaman ini dimasukkan, gambar akan memperlihatkan tabel produk lengkap beserta tombol modal pengelolaan data.",
    ),
    (
        "Halaman Manajemen Pesanan",
        "Halaman Manajemen Pesanan telah diimplementasikan sebagai daftar transaksi pelanggan yang dapat dipantau admin secara menyeluruh. Halaman ini menampilkan filter status pesanan, tabel berisi invoice, identitas pelanggan, tanggal transaksi, total pembayaran, status pesanan, tipe pengiriman, serta tombol detail dan pembatalan pesanan. Admin dapat menyaring data berdasarkan status agar pemantauan pekerjaan menjadi lebih fokus, kemudian membuka salah satu transaksi untuk pemeriksaan lanjutan. Alur penggunaan dimulai dari pemilihan filter, peninjauan daftar pesanan, lalu pembukaan detail atau pembatalan pesanan jika masih memenuhi syarat operasional. Apabila screenshot halaman ini ditampilkan, gambar akan menunjukkan form filter di bagian atas dan tabel daftar pesanan pada bagian utama halaman.",
    ),
    (
        "Halaman Detail Pesanan dan Perubahan Status",
        "Halaman Detail Pesanan dan Perubahan Status telah direalisasikan untuk memusatkan seluruh proses pengendalian transaksi admin. Komponen yang tampil meliputi identitas pesanan, data akun pelanggan, snapshot data pengiriman saat checkout, tabel item pesanan, riwayat pembayaran, bukti transfer yang dapat diperbesar, serta panel aksi admin untuk mengubah status pesanan, membatalkan pesanan, memverifikasi pembayaran, menolak pembayaran, dan mengisi kode resi ketika status berubah menjadi dikirim. Halaman ini memungkinkan admin menjalankan alur pemrosesan dari tahap menunggu pembayaran hingga selesai dalam satu tampilan. Alur penggunaan dimulai dari pembukaan invoice tertentu, peninjauan data, lalu pengambilan keputusan operasional sesuai kondisi transaksi. Apabila screenshot halaman ini dimasukkan, gambar akan memperlihatkan detail transaksi lengkap dengan panel aksi status di sisi samping.",
    ),
    (
        "Halaman Verifikasi Pembayaran",
        "Halaman Verifikasi Pembayaran telah diimplementasikan sebagai daftar khusus untuk seluruh bukti pembayaran berstatus pending. Halaman ini menampilkan data pelanggan, invoice, bank pengirim dan tujuan, nominal transfer, total pesanan, waktu unggah, tombol lihat bukti, serta tombol setujui dan tolak. Admin dapat membandingkan nominal transfer dengan nilai pesanan, memeriksa kejelasan bukti gambar, lalu memutuskan hasil verifikasi. Pada implementasi aktual, pembayaran yang disetujui akan membuat pesanan berlanjut ke tahap diproses, sedangkan pembayaran yang ditolak akan diberi alasan dan pesanan kembali ke tahap menunggu bayar agar pelanggan dapat mengunggah bukti yang benar. Apabila screenshot halaman ini ditampilkan, gambar akan memperlihatkan tabel pembayaran pending lengkap dengan tombol verifikasi pada setiap baris transaksi.",
    ),
    (
        "Halaman Laporan Penjualan",
        "Halaman Laporan Penjualan telah diimplementasikan untuk menyajikan hasil penjualan berdasarkan periode tanggal tertentu. Halaman ini memuat filter tanggal mulai dan akhir, tombol ekspor CSV, kartu ringkasan total pesanan dan total pendapatan, tabel daftar pesanan, serta tabel produk terlaris. Admin dapat melakukan penyaringan periode, melihat akumulasi transaksi berstatus selesai, dan mengunduh data penjualan untuk keperluan dokumentasi atau analisis. Alur penggunaan dimulai dari penentuan rentang tanggal, dilanjutkan dengan pembacaan ringkasan dan detail tabel penjualan. Apabila screenshot halaman ini dimasukkan, gambar akan menunjukkan panel filter, kartu ringkasan penjualan, serta tabel daftar pesanan dan produk terlaris pada periode terpilih.",
    ),
    (
        "Halaman Rekap Stok",
        "Halaman Rekap Stok telah diimplementasikan untuk menampilkan kondisi persediaan produk secara menyeluruh. Komponen utama halaman ini terdiri atas kartu total nilai stok, jumlah stok rendah, jumlah stok habis, serta tabel detail produk yang memuat kategori, stok tersedia, stok dibooking, harga satuan, nilai stok, dan status persediaan. Admin dapat menggunakan halaman ini untuk mengidentifikasi produk yang harus segera diprioritaskan dalam pengelolaan stok. Alur penggunaan dimulai dari pembacaan indikator ringkas, kemudian admin menelusuri tabel detail produk untuk mengetahui posisi stok setiap item. Apabila screenshot halaman ini ditampilkan, gambar akan memperlihatkan kartu rekap stok pada bagian atas dan tabel persediaan produk pada bagian bawah.",
    ),
    (
        "Halaman Data Pelanggan",
        "Halaman Data Pelanggan telah diimplementasikan untuk menyajikan rekap pelanggan beserta statistik transaksinya. Halaman ini menampilkan kartu ringkasan total pelanggan, pelanggan aktif, total pesanan, dan total pendapatan pelanggan, kemudian dilengkapi kolom pencarian serta tabel data pelanggan yang diurutkan berdasarkan total belanja tertinggi. Admin dapat mencari pelanggan tertentu, melihat jumlah pesanan yang pernah dilakukan, dan meninjau besaran total belanja masing-masing pelanggan. Alur penggunaan dimulai dari membaca ringkasan umum, menggunakan fitur pencarian bila diperlukan, lalu menelaah tabel pelanggan secara rinci. Apabila screenshot halaman ini dimasukkan, gambar akan menunjukkan kartu statistik pelanggan di atas dan tabel data pelanggan dengan kolom pencarian pada area utama.",
    ),
]


def build_document():
    doc = Document()
    configure_document(doc)

    add_center_heading(doc, "BAB IV")
    add_center_heading(doc, "HASIL DAN PEMBAHASAN")

    add_left_heading(doc, "4.4 Perancangan Antarmuka")
    add_left_heading(doc, "4.4.1 Halaman Pelanggan")
    for idx, (title, body) in enumerate(customer_44):
        add_item(doc, chr(ord("a") + idx), title, body)

    add_left_heading(doc, "4.4.2 Halaman Admin")
    for idx, (title, body) in enumerate(admin_44):
        add_item(doc, chr(ord("a") + idx), title, body)

    add_left_heading(doc, "4.5 Implementasi Sistem")
    add_left_heading(doc, "4.5.1 Halaman Pelanggan")
    for idx, (title, body) in enumerate(customer_45):
        add_item(doc, chr(ord("a") + idx), title, body)

    add_left_heading(doc, "4.5.2 Halaman Admin")
    for idx, (title, body) in enumerate(admin_45):
        add_item(doc, chr(ord("a") + idx), title, body)

    doc.save(OUTPUT_FILE)


if __name__ == "__main__":
    build_document()
    print(f"Generated {OUTPUT_FILE}")
