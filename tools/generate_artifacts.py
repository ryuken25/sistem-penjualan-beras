from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent.parent
DIAGRAM_DIR = ROOT / 'diagram'
DOCS_DIR = ROOT / 'docs'
AUDIT_DIR = ROOT / 'audit'
SCREENSHOT_DIR = ROOT / 'screenshots'
OUTPUT_AUDIT_DIR = ROOT / 'output_audit'
OUTPUT_DIAGRAM_DIR = ROOT / 'output_diagram'
OUTPUT_DOCX_DIR = ROOT / 'output_docx'


@dataclass
class Node:
    node_id: str
    label: str
    x: int
    y: int
    width: int
    height: int
    kind: str


@dataclass
class Edge:
    edge_id: str
    source: str
    target: str
    label: str


SYSTEM_PROFILE = {
    'name': 'Sistem Informasi Penjualan Beras Berbasis Website',
    'organization': 'UD Tulus Sari Merta',
    'domain': 'sistem informasi penjualan beras berbasis web untuk administrasi penjualan internal',
    'stack': [
        'PHP 8.2',
        'CodeIgniter 4.7',
        'MySQL/MariaDB',
        'Bootstrap 5',
        'Chart.js',
    ],
    'roles': ['admin', 'pegawai'],
    'scope_notes': [
        'Tidak ada portal pelanggan publik.',
        'Tidak ada payment gateway.',
        'Tidak ada manajemen stok, produksi, atau penggilingan.',
        'Tidak ada multi cabang dan multi gudang.',
    ],
}


FEATURE_AUDIT = {
    'public_pages': ['login'],
    'auth': ['login', 'logout', 'session-based authentication'],
    'dashboards': ['dashboard admin', 'dashboard pegawai'],
    'master_modules': [
        'kelola pengguna',
        'kelola produk beras',
        'kelola harga beras',
        'kelola template transaksi cepat',
        'kelola mode pembatasan penjualan',
    ],
    'transaction_modules': [
        'transaksi manual berdasarkan qty 5 kg, 10 kg, dan 25 kg',
        'transaksi berbasis template cepat',
        'riwayat transaksi',
    ],
    'reporting_modules': [
        'laporan penjualan dengan filter tanggal, produk, dan pencatat',
        'grafik penjualan harian dan bulanan',
    ],
    'profile_modules': [
        'ubah nama lengkap',
        'ubah username',
        'upload foto profil',
        'ubah password',
    ],
}


TABLES = [
    {
        'name': 'users',
        'code': 'D1',
        'description': 'menyimpan akun admin dan pegawai',
        'fields': [
            'id (PK)', 'full_name', 'username', 'password_hash', 'profile_photo', 'role', 'is_active',
            'created_at', 'updated_at', 'deleted_at'
        ],
        'source': 'migration + model + controller',
    },
    {
        'name': 'products',
        'code': 'D2',
        'description': 'menyimpan master produk beras',
        'fields': ['id (PK)', 'product_code', 'product_name', 'weight_kg', 'is_active', 'created_at', 'updated_at', 'deleted_at'],
        'source': 'migration + model + controller',
    },
    {
        'name': 'product_prices',
        'code': 'D3',
        'description': 'menyimpan histori dan harga aktif produk',
        'fields': ['id (PK)', 'product_id (FK)', 'price', 'effective_date', 'is_current', 'updated_by (FK)', 'created_at', 'updated_at'],
        'source': 'migration + model + controller',
    },
    {
        'name': 'quick_templates',
        'code': 'D4',
        'description': 'menyimpan header template transaksi cepat',
        'fields': ['id (PK)', 'template_code', 'template_name', 'qty_5kg', 'qty_10kg', 'qty_25kg', 'is_active', 'created_by (FK)', 'created_at', 'updated_at', 'deleted_at'],
        'source': 'migration + model + controller',
    },
    {
        'name': 'quick_template_items',
        'code': 'D5',
        'description': 'menyimpan detail produk per template',
        'fields': ['id (PK)', 'template_id (FK)', 'product_id (FK)', 'quantity'],
        'source': 'migration + model + controller',
    },
    {
        'name': 'sale_limit_settings',
        'code': 'D6',
        'description': 'menyimpan konfigurasi mode pembatasan penjualan',
        'fields': ['id (PK)', 'is_enabled', 'max_total_kg', 'updated_by (FK)', 'created_at', 'updated_at'],
        'source': 'migration + model + controller',
    },
    {
        'name': 'sales_transactions',
        'code': 'D7',
        'description': 'menyimpan header transaksi penjualan',
        'fields': [
            'id (PK)', 'invoice_number', 'transaction_date', 'created_by (FK)', 'template_id (FK, nullable)',
            'customer_name', 'qty_5kg', 'qty_10kg', 'qty_25kg', 'price_5kg', 'price_10kg', 'price_25kg',
            'subtotal_5kg', 'subtotal_10kg', 'subtotal_25kg', 'total_items', 'total_kg', 'total_harga',
            'grand_total', 'source_transaksi', 'notes', 'created_at', 'updated_at'
        ],
        'source': 'migration + service + model + controller',
    },
    {
        'name': 'sales_transaction_items',
        'code': 'D8',
        'description': 'menyimpan detail item transaksi dan snapshot harga',
        'fields': ['id (PK)', 'transaction_id (FK)', 'product_id (FK, nullable)', 'product_name_snapshot', 'weight_kg_snapshot', 'unit_price_snapshot', 'quantity', 'subtotal', 'total_kg_item'],
        'source': 'migration + service + model + controller',
    },
]


RELATIONSHIPS = [
    ('users', 'product_prices', '1:N', 'updated_by'),
    ('users', 'quick_templates', '1:N', 'created_by'),
    ('users', 'sale_limit_settings', '1:N', 'updated_by'),
    ('users', 'sales_transactions', '1:N', 'created_by'),
    ('products', 'product_prices', '1:N', 'product_id'),
    ('products', 'quick_template_items', '1:N', 'product_id'),
    ('products', 'sales_transaction_items', '1:N', 'product_id'),
    ('quick_templates', 'quick_template_items', '1:N', 'template_id'),
    ('quick_templates', 'sales_transactions', '1:N', 'template_id'),
    ('sales_transactions', 'sales_transaction_items', '1:N', 'transaction_id'),
]


BALANCING = [
    ('D1 users', 'data_profil_akun (×1), data_pengguna (×1)', 'info_akun (×1), info_pengguna (×1)', 'Seimbang'),
    ('D2 products', 'data_produk (×1)', 'info_produk (×1)', 'Seimbang'),
    ('D3 product_prices', 'data_harga_produk (×1)', 'info_harga_produk (×1)', 'Seimbang'),
    ('D4 quick_templates', 'data_template (×1)', 'info_template (×1)', 'Seimbang'),
    ('D5 quick_template_items', 'data_item_template (×1)', 'info_item_template (×1)', 'Seimbang'),
    ('D6 sale_limit_settings', 'data_limit_penjualan (×1)', 'info_limit_penjualan (×1)', 'Seimbang'),
    ('D7 sales_transactions', 'data_transaksi (×1)', 'info_transaksi (×1)', 'Seimbang'),
    ('D8 sales_transaction_items', 'data_detail_transaksi (×1)', 'info_detail_transaksi (×1)', 'Seimbang'),
]


SCREENSHOT_PLAN = [
    ('01_Login.png', 'Halaman login sistem.'),
    ('02_Dashboard_Admin.png', 'Dashboard admin yang menampilkan ringkasan transaksi, status limit, dan grafik.'),
    ('03_Kelola_Pengguna.png', 'Halaman kelola pengguna.'),
    ('04_Kelola_Produk.png', 'Halaman kelola produk beras.'),
    ('05_Kelola_Harga.png', 'Halaman kelola harga dan histori harga.'),
    ('06_Template_Transaksi.png', 'Halaman kelola template transaksi cepat.'),
    ('07_Mode_Limit.png', 'Halaman pengaturan mode pembatasan penjualan.'),
    ('08_Laporan_Penjualan.png', 'Halaman laporan penjualan dengan filter dan tabel hasil.'),
    ('09_Grafik_Penjualan.png', 'Halaman grafik penjualan.'),
    ('10_Profil_Admin.png', 'Halaman profil admin.'),
    ('11_Dashboard_Pegawai.png', 'Dashboard pegawai.'),
    ('12_Transaksi_Manual.png', 'Halaman input transaksi manual.'),
    ('13_Transaksi_Template.png', 'Halaman transaksi berbasis template cepat.'),
    ('14_Riwayat_Transaksi.png', 'Halaman riwayat transaksi.'),
]


def make_context_diagram() -> dict:
    return {
        'title': 'Diagram Konteks',
        'nodes': [
            Node('admin', 'Admin', 60, 180, 130, 60, 'external'),
            Node('system', 'Sistem Informasi\nPenjualan Beras', 320, 150, 240, 100, 'process'),
            Node('pegawai', 'Pegawai', 690, 180, 130, 60, 'external'),
        ],
        'edges': [
            Edge('e1', 'admin', 'system', 'data_login\ndata_profil\ndata_pengguna\ndata_produk_harga\ndata_template_limit\ndata_permintaan_laporan'),
            Edge('e2', 'system', 'admin', 'info_autentikasi\ninfo_dashboard_admin\ninfo_pengguna\ninfo_produk_harga\ninfo_template_limit\ninfo_transaksi\ninfo_laporan_grafik'),
            Edge('e3', 'pegawai', 'system', 'data_login\ndata_profil\ndata_transaksi\ndata_pilih_template\ndata_permintaan_dashboard'),
            Edge('e4', 'system', 'pegawai', 'info_autentikasi\ninfo_dashboard_pegawai\ninfo_profil\ninfo_template_transaksi\ninfo_transaksi'),
        ],
    }


def make_level0_diagram() -> dict:
    return {
        'title': 'DFD Level 0',
        'nodes': [
            Node('admin', 'Admin', 30, 170, 120, 55, 'external'),
            Node('pegawai', 'Pegawai', 30, 360, 120, 55, 'external'),
            Node('p1', '1.0 Autentikasi\ndan Profil', 220, 40, 170, 80, 'process'),
            Node('p2', '2.0 Kelola\nPengguna', 220, 150, 170, 80, 'process'),
            Node('p3', '3.0 Kelola\nProduk & Harga', 220, 260, 170, 80, 'process'),
            Node('p4', '4.0 Kelola Template\n& Mode Limit', 220, 370, 170, 80, 'process'),
            Node('p5', '5.0 Proses\nTransaksi Penjualan', 520, 150, 190, 85, 'process'),
            Node('p6', '6.0 Dashboard,\nLaporan & Grafik', 520, 320, 190, 85, 'process'),
            Node('d1', 'D1 users', 800, 30, 170, 55, 'store'),
            Node('d2', 'D2 products', 800, 100, 170, 55, 'store'),
            Node('d3', 'D3 product_prices', 800, 170, 170, 55, 'store'),
            Node('d4', 'D4 quick_templates', 800, 240, 170, 55, 'store'),
            Node('d5', 'D5 quick_template_items', 800, 310, 170, 55, 'store'),
            Node('d6', 'D6 sale_limit_settings', 800, 380, 170, 55, 'store'),
            Node('d7', 'D7 sales_transactions', 800, 450, 170, 55, 'store'),
            Node('d8', 'D8 sales_transaction_items', 800, 520, 170, 55, 'store'),
        ],
        'edges': [
            Edge('e1', 'admin', 'p1', 'data_login_profil'),
            Edge('e2', 'pegawai', 'p1', 'data_login_profil'),
            Edge('e3', 'p1', 'admin', 'info_autentikasi_profil'),
            Edge('e4', 'p1', 'pegawai', 'info_autentikasi_profil'),
            Edge('e5', 'admin', 'p2', 'data_pengguna'),
            Edge('e6', 'p2', 'admin', 'info_pengguna'),
            Edge('e7', 'admin', 'p3', 'data_produk_harga'),
            Edge('e8', 'p3', 'admin', 'info_produk_harga'),
            Edge('e9', 'admin', 'p4', 'data_template_limit'),
            Edge('e10', 'p4', 'admin', 'info_template_limit'),
            Edge('e11', 'admin', 'p5', 'data_transaksi'),
            Edge('e12', 'pegawai', 'p5', 'data_transaksi'),
            Edge('e13', 'p5', 'admin', 'info_transaksi'),
            Edge('e14', 'p5', 'pegawai', 'info_transaksi'),
            Edge('e15', 'admin', 'p6', 'data_permintaan_dashboard_laporan'),
            Edge('e16', 'pegawai', 'p6', 'data_permintaan_dashboard'),
            Edge('e17', 'p6', 'admin', 'info_dashboard_laporan_grafik'),
            Edge('e18', 'p6', 'pegawai', 'info_dashboard_pegawai'),
            Edge('e19', 'p1', 'd1', 'data_profil_akun'),
            Edge('e20', 'd1', 'p1', 'info_akun'),
            Edge('e21', 'p2', 'd1', 'data_pengguna'),
            Edge('e22', 'd1', 'p2', 'info_pengguna'),
            Edge('e23', 'p3', 'd2', 'data_produk'),
            Edge('e24', 'd2', 'p3', 'info_produk'),
            Edge('e25', 'p3', 'd3', 'data_harga_produk'),
            Edge('e26', 'd3', 'p3', 'info_harga_produk'),
            Edge('e27', 'p4', 'd4', 'data_template'),
            Edge('e28', 'd4', 'p4', 'info_template'),
            Edge('e29', 'p4', 'd5', 'data_item_template'),
            Edge('e30', 'd5', 'p4', 'info_item_template'),
            Edge('e31', 'p4', 'd6', 'data_limit_penjualan'),
            Edge('e32', 'd6', 'p4', 'info_limit_penjualan'),
            Edge('e33', 'p5', 'd7', 'data_transaksi'),
            Edge('e34', 'd7', 'p5', 'info_transaksi'),
            Edge('e35', 'p5', 'd8', 'data_detail_transaksi'),
            Edge('e36', 'd8', 'p5', 'info_detail_transaksi'),
            Edge('e37', 'p2', 'p6', 'info_pengguna_aktif'),
            Edge('e38', 'p3', 'p5', 'info_produk_harga_aktif'),
            Edge('e39', 'p3', 'p6', 'info_produk_harga_aktif'),
            Edge('e40', 'p4', 'p5', 'info_template_limit_aktif'),
            Edge('e41', 'p4', 'p6', 'info_limit_penjualan'),
            Edge('e42', 'p5', 'p6', 'info_transaksi_tercatat'),
        ],
    }


def make_level1_diagrams() -> list[dict]:
    return [
        {
            'title': 'DFD Level 1 P1',
            'filename': 'DFD_Level1_P1',
            'nodes': [
                Node('admin', 'Admin', 40, 80, 120, 55, 'external'),
                Node('pegawai', 'Pegawai', 40, 210, 120, 55, 'external'),
                Node('p11', '1.1 Validasi\nLogin', 260, 80, 170, 75, 'process'),
                Node('p12', '1.2 Perbarui\nProfil', 260, 220, 170, 75, 'process'),
                Node('d1', 'D1 users', 560, 145, 170, 55, 'store'),
            ],
            'edges': [
                Edge('e1', 'admin', 'p11', 'data_login'),
                Edge('e2', 'pegawai', 'p11', 'data_login'),
                Edge('e3', 'd1', 'p11', 'info_akun'),
                Edge('e4', 'p11', 'admin', 'info_autentikasi'),
                Edge('e5', 'p11', 'pegawai', 'info_autentikasi'),
                Edge('e6', 'admin', 'p12', 'data_profil'),
                Edge('e7', 'pegawai', 'p12', 'data_profil'),
                Edge('e8', 'p12', 'd1', 'data_profil_akun'),
                Edge('e9', 'p12', 'admin', 'info_profil'),
                Edge('e10', 'p12', 'pegawai', 'info_profil'),
            ],
        },
        {
            'title': 'DFD Level 1 P2',
            'filename': 'DFD_Level1_P2',
            'nodes': [
                Node('admin', 'Admin', 40, 140, 120, 55, 'external'),
                Node('p21', '2.1 Simpan\nData Pengguna', 240, 80, 180, 75, 'process'),
                Node('p22', '2.2 Tampilkan\nData Pengguna', 240, 220, 180, 75, 'process'),
                Node('d1', 'D1 users', 560, 145, 170, 55, 'store'),
            ],
            'edges': [
                Edge('e1', 'admin', 'p21', 'data_pengguna'),
                Edge('e2', 'p21', 'd1', 'data_pengguna'),
                Edge('e3', 'd1', 'p22', 'info_pengguna'),
                Edge('e4', 'p22', 'admin', 'info_pengguna'),
            ],
        },
        {
            'title': 'DFD Level 1 P3',
            'filename': 'DFD_Level1_P3',
            'nodes': [
                Node('admin', 'Admin', 40, 150, 120, 55, 'external'),
                Node('p31', '3.1 Simpan\nData Produk', 240, 60, 180, 75, 'process'),
                Node('p32', '3.2 Sajikan\nInfo Produk', 240, 150, 180, 75, 'process'),
                Node('p33', '3.3 Simpan\nHarga Produk', 240, 240, 180, 75, 'process'),
                Node('p34', '3.4 Sajikan\nInfo Harga', 240, 330, 180, 75, 'process'),
                Node('d2', 'D2 products', 560, 110, 170, 55, 'store'),
                Node('d3', 'D3 product_prices', 560, 280, 170, 55, 'store'),
            ],
            'edges': [
                Edge('e1', 'admin', 'p31', 'data_produk'),
                Edge('e2', 'p31', 'd2', 'data_produk'),
                Edge('e3', 'd2', 'p32', 'info_produk'),
                Edge('e4', 'p32', 'admin', 'info_produk'),
                Edge('e5', 'admin', 'p33', 'data_harga_produk'),
                Edge('e6', 'p33', 'd3', 'data_harga_produk'),
                Edge('e7', 'd3', 'p34', 'info_harga_produk'),
                Edge('e8', 'p34', 'admin', 'info_harga_produk'),
            ],
        },
        {
            'title': 'DFD Level 1 P4',
            'filename': 'DFD_Level1_P4',
            'nodes': [
                Node('admin', 'Admin', 30, 180, 120, 55, 'external'),
                Node('p41', '4.1 Simpan\nTemplate', 220, 40, 170, 75, 'process'),
                Node('p42', '4.2 Sajikan\nTemplate', 220, 130, 170, 75, 'process'),
                Node('p43', '4.3 Simpan\nLimit Penjualan', 220, 250, 170, 75, 'process'),
                Node('p44', '4.4 Sajikan\nLimit Penjualan', 220, 340, 170, 75, 'process'),
                Node('d4', 'D4 quick_templates', 540, 30, 180, 55, 'store'),
                Node('d5', 'D5 quick_template_items', 540, 130, 180, 55, 'store'),
                Node('d6', 'D6 sale_limit_settings', 540, 300, 180, 55, 'store'),
            ],
            'edges': [
                Edge('e1', 'admin', 'p41', 'data_template'),
                Edge('e2', 'p41', 'd4', 'data_template'),
                Edge('e3', 'p41', 'd5', 'data_item_template'),
                Edge('e4', 'd4', 'p42', 'info_template'),
                Edge('e5', 'd5', 'p42', 'info_item_template'),
                Edge('e6', 'p42', 'admin', 'info_template'),
                Edge('e7', 'admin', 'p43', 'data_limit_penjualan'),
                Edge('e8', 'p43', 'd6', 'data_limit_penjualan'),
                Edge('e9', 'd6', 'p44', 'info_limit_penjualan'),
                Edge('e10', 'p44', 'admin', 'info_limit_penjualan'),
            ],
        },
        {
            'title': 'DFD Level 1 P5',
            'filename': 'DFD_Level1_P5',
            'nodes': [
                Node('admin', 'Admin', 30, 80, 120, 55, 'external'),
                Node('pegawai', 'Pegawai', 30, 260, 120, 55, 'external'),
                Node('p51', '5.1 Terima\nData Transaksi', 230, 60, 180, 75, 'process'),
                Node('p52', '5.2 Hitung &\nValidasi Transaksi', 230, 180, 180, 75, 'process'),
                Node('p53', '5.3 Simpan &\nSajikan Transaksi', 230, 300, 180, 75, 'process'),
                Node('d7', 'D7 sales_transactions', 550, 130, 180, 55, 'store'),
                Node('d8', 'D8 sales_transaction_items', 550, 290, 180, 55, 'store'),
            ],
            'edges': [
                Edge('e1', 'admin', 'p51', 'data_transaksi'),
                Edge('e2', 'pegawai', 'p51', 'data_transaksi'),
                Edge('e3', 'p51', 'p52', 'data_transaksi_terverifikasi'),
                Edge('e4', 'p52', 'p53', 'data_transaksi_final'),
                Edge('e5', 'p53', 'd7', 'data_transaksi'),
                Edge('e6', 'd7', 'p53', 'info_transaksi'),
                Edge('e7', 'p53', 'd8', 'data_detail_transaksi'),
                Edge('e8', 'd8', 'p53', 'info_detail_transaksi'),
                Edge('e9', 'p53', 'admin', 'info_transaksi'),
                Edge('e10', 'p53', 'pegawai', 'info_transaksi'),
            ],
        },
        {
            'title': 'DFD Level 1 P6',
            'filename': 'DFD_Level1_P6',
            'nodes': [
                Node('admin', 'Admin', 30, 80, 120, 55, 'external'),
                Node('pegawai', 'Pegawai', 30, 280, 120, 55, 'external'),
                Node('p61', '6.1 Susun\nLaporan', 250, 50, 180, 75, 'process'),
                Node('p62', '6.2 Bentuk\nGrafik', 250, 160, 180, 75, 'process'),
                Node('p63', '6.3 Susun\nDashboard', 250, 270, 180, 75, 'process'),
                Node('p2', '2.0 Kelola\nPengguna', 540, 20, 170, 60, 'process_link'),
                Node('p3', '3.0 Kelola\nProduk & Harga', 540, 130, 170, 60, 'process_link'),
                Node('p4', '4.0 Kelola Template\n& Mode Limit', 540, 240, 170, 60, 'process_link'),
                Node('p5', '5.0 Proses\nTransaksi Penjualan', 540, 350, 170, 60, 'process_link'),
            ],
            'edges': [
                Edge('e1', 'admin', 'p61', 'data_filter_laporan'),
                Edge('e2', 'admin', 'p62', 'data_permintaan_grafik'),
                Edge('e3', 'pegawai', 'p63', 'data_permintaan_dashboard'),
                Edge('e4', 'admin', 'p63', 'data_permintaan_dashboard'),
                Edge('e5', 'p2', 'p61', 'info_pengguna_aktif'),
                Edge('e6', 'p3', 'p61', 'info_produk_harga_aktif'),
                Edge('e7', 'p5', 'p61', 'info_transaksi_tercatat'),
                Edge('e8', 'p5', 'p62', 'info_transaksi_tercatat'),
                Edge('e9', 'p4', 'p63', 'info_limit_penjualan'),
                Edge('e10', 'p5', 'p63', 'info_transaksi_tercatat'),
                Edge('e11', 'p61', 'admin', 'info_laporan_penjualan'),
                Edge('e12', 'p62', 'admin', 'info_grafik_penjualan'),
                Edge('e13', 'p63', 'admin', 'info_dashboard_admin'),
                Edge('e14', 'p63', 'pegawai', 'info_dashboard_pegawai'),
            ],
        },
    ]


def make_erd_chen() -> dict:
    return {
        'title': 'ERD Chen',
        'nodes': [
            Node('users', 'users', 60, 40, 140, 55, 'entity'),
            Node('products', 'products', 60, 300, 140, 55, 'entity'),
            Node('prices', 'product_prices', 360, 40, 170, 55, 'entity'),
            Node('qt', 'quick_templates', 360, 170, 170, 55, 'entity'),
            Node('qti', 'quick_template_items', 660, 170, 190, 55, 'entity'),
            Node('limit', 'sale_limit_settings', 360, 300, 180, 55, 'entity'),
            Node('st', 'sales_transactions', 660, 40, 190, 55, 'entity'),
            Node('sti', 'sales_transaction_items', 930, 40, 210, 55, 'entity'),
            Node('r1', 'memperbarui', 250, 40, 95, 55, 'relationship'),
            Node('r2', 'membuat', 250, 170, 95, 55, 'relationship'),
            Node('r3', 'mengatur', 250, 300, 95, 55, 'relationship'),
            Node('r4', 'berisi', 560, 170, 95, 55, 'relationship'),
            Node('r5', 'mencatat', 560, 40, 95, 55, 'relationship'),
            Node('r6', 'terdiri dari', 860, 40, 110, 55, 'relationship'),
            Node('r7', 'mereferensi', 930, 170, 110, 55, 'relationship'),
        ],
        'edges': [
            Edge('e1', 'users', 'r1', '1'), Edge('e2', 'r1', 'prices', 'N'),
            Edge('e3', 'users', 'r2', '1'), Edge('e4', 'r2', 'qt', 'N'),
            Edge('e5', 'users', 'r3', '1'), Edge('e6', 'r3', 'limit', 'N'),
            Edge('e7', 'qt', 'r4', '1'), Edge('e8', 'r4', 'qti', 'N'),
            Edge('e9', 'users', 'r5', '1'), Edge('e10', 'r5', 'st', 'N'),
            Edge('e11', 'st', 'r6', '1'), Edge('e12', 'r6', 'sti', 'N'),
            Edge('e13', 'products', 'r7', '1'), Edge('e14', 'r7', 'sti', 'N'),
            Edge('e15', 'products', 'prices', '1:N'), Edge('e16', 'products', 'qti', '1:N'), Edge('e17', 'qt', 'st', '1:N'),
        ],
    }


def make_erd_crowsfoot() -> dict:
    nodes = [
        Node('users', 'users\nPK id\nfull_name\nusername\nrole\nis_active', 40, 30, 180, 120, 'table'),
        Node('products', 'products\nPK id\nproduct_code\nproduct_name\nweight_kg\nis_active', 40, 250, 180, 120, 'table'),
        Node('prices', 'product_prices\nPK id\nFK product_id\nprice\neffective_date\nis_current\nFK updated_by', 310, 30, 200, 140, 'table'),
        Node('qt', 'quick_templates\nPK id\ntemplate_code\ntemplate_name\nqty_5kg\nqty_10kg\nqty_25kg\nFK created_by', 310, 220, 200, 165, 'table'),
        Node('qti', 'quick_template_items\nPK id\nFK template_id\nFK product_id\nquantity', 600, 220, 200, 120, 'table'),
        Node('limit', 'sale_limit_settings\nPK id\nis_enabled\nmax_total_kg\nFK updated_by', 310, 430, 200, 110, 'table'),
        Node('st', 'sales_transactions\nPK id\ninvoice_number\ntransaction_date\nFK created_by\nFK template_id\ntotal_items\ntotal_kg\ngrand_total', 600, 20, 215, 170, 'table'),
        Node('sti', 'sales_transaction_items\nPK id\nFK transaction_id\nFK product_id\nquantity\nsubtotal\ntotal_kg_item', 900, 40, 210, 130, 'table'),
    ]
    return {
        'title': 'ERD Crow\'s Foot',
        'nodes': nodes,
        'edges': [
            Edge('e1', 'users', 'prices', '1:N updated_by'),
            Edge('e2', 'users', 'qt', '1:N created_by'),
            Edge('e3', 'users', 'limit', '1:N updated_by'),
            Edge('e4', 'users', 'st', '1:N created_by'),
            Edge('e5', 'products', 'prices', '1:N product_id'),
            Edge('e6', 'products', 'qti', '1:N product_id'),
            Edge('e7', 'products', 'sti', '1:N product_id'),
            Edge('e8', 'qt', 'qti', '1:N template_id'),
            Edge('e9', 'qt', 'st', '1:N template_id'),
            Edge('e10', 'st', 'sti', '1:N transaction_id'),
        ],
    }


def make_conceptual_db() -> dict:
    return {
        'title': 'Konseptual Basis Data',
        'nodes': [
            Node('master', 'Kelompok Master\nusers\nproducts\nproduct_prices', 80, 120, 220, 140, 'group'),
            Node('config', 'Kelompok Konfigurasi\nquick_templates\nquick_template_items\nsale_limit_settings', 390, 120, 260, 160, 'group'),
            Node('trx', 'Kelompok Transaksi\nsales_transactions\nsales_transaction_items', 760, 120, 250, 130, 'group'),
        ],
        'edges': [
            Edge('e1', 'master', 'config', 'info_master_produk_pengguna'),
            Edge('e2', 'master', 'trx', 'info_master_produk_pengguna'),
            Edge('e3', 'config', 'trx', 'info_template_dan_limit'),
        ],
    }


def make_table_structure_diagram() -> dict:
    return {
        'title': 'Struktur Tabel Relasi',
        'nodes': [
            Node('u', 'users\nPK id\nfull_name\nusername\npassword_hash\nprofile_photo\nrole\nis_active', 30, 20, 200, 170, 'table'),
            Node('p', 'products\nPK id\nproduct_code\nproduct_name\nweight_kg\nis_active', 30, 240, 200, 150, 'table'),
            Node('pp', 'product_prices\nPK id\nFK product_id\nprice\neffective_date\nis_current\nFK updated_by', 280, 20, 220, 170, 'table'),
            Node('qt', 'quick_templates\nPK id\ntemplate_code\ntemplate_name\nqty_5kg\nqty_10kg\nqty_25kg\nis_active\nFK created_by', 280, 220, 220, 185, 'table'),
            Node('qti', 'quick_template_items\nPK id\nFK template_id\nFK product_id\nquantity', 560, 250, 220, 125, 'table'),
            Node('sl', 'sale_limit_settings\nPK id\nis_enabled\nmax_total_kg\nFK updated_by', 280, 450, 220, 125, 'table'),
            Node('st', 'sales_transactions\nPK id\ninvoice_number\ntransaction_date\nFK created_by\nFK template_id\nqty_5kg\nqty_10kg\nqty_25kg\nprice_5kg\nprice_10kg\nprice_25kg\nsubtotal_5kg\nsubtotal_10kg\nsubtotal_25kg\ntotal_items\ntotal_kg\ntotal_harga\ngrand_total\nsource_transaksi\nnotes', 820, 20, 250, 330, 'table'),
            Node('sti', 'sales_transaction_items\nPK id\nFK transaction_id\nFK product_id\nproduct_name_snapshot\nweight_kg_snapshot\nunit_price_snapshot\nquantity\nsubtotal\ntotal_kg_item', 1120, 80, 250, 230, 'table'),
        ],
        'edges': [
            Edge('e1', 'u', 'pp', '1:N updated_by'),
            Edge('e2', 'u', 'qt', '1:N created_by'),
            Edge('e3', 'u', 'sl', '1:N updated_by'),
            Edge('e4', 'u', 'st', '1:N created_by'),
            Edge('e5', 'p', 'pp', '1:N product_id'),
            Edge('e6', 'p', 'qti', '1:N product_id'),
            Edge('e7', 'p', 'sti', '1:N product_id'),
            Edge('e8', 'qt', 'qti', '1:N template_id'),
            Edge('e9', 'qt', 'st', '1:N template_id'),
            Edge('e10', 'st', 'sti', '1:N transaction_id'),
        ],
    }


ALL_DIAGRAMS = [
    ('Diagram_Konteks', make_context_diagram()),
    ('DFD_Level0', make_level0_diagram()),
    *[(d['filename'], d) for d in make_level1_diagrams()],
    ('ERD_Chen', make_erd_chen()),
    ('ERD_CrowsFoot', make_erd_crowsfoot()),
    ('Konseptual_Basis_Data', make_conceptual_db()),
    ('Struktur_Tabel_Relasi', make_table_structure_diagram()),
]


def style_for_kind(kind: str) -> dict:
    styles = {
        'external': {'shape': 'rectangle', 'fill': '#F3F4F6', 'border': '#374151'},
        'process': {'shape': 'ellipse', 'fill': '#DBEAFE', 'border': '#1D4ED8'},
        'process_link': {'shape': 'ellipse', 'fill': '#E0E7FF', 'border': '#4338CA'},
        'store': {'shape': 'roundrectangle', 'fill': '#DCFCE7', 'border': '#15803D'},
        'entity': {'shape': 'rectangle', 'fill': '#FEF3C7', 'border': '#D97706'},
        'relationship': {'shape': 'diamond', 'fill': '#FCE7F3', 'border': '#BE185D'},
        'table': {'shape': 'rectangle', 'fill': '#F8FAFC', 'border': '#0F172A'},
        'group': {'shape': 'roundrectangle', 'fill': '#F5F3FF', 'border': '#7C3AED'},
    }
    return styles[kind]


def graphml_node(node: Node) -> str:
    style = style_for_kind(node.kind)
    return f'''    <node id="{escape(node.node_id)}">
      <data key="d0">
        <y:ShapeNode>
          <y:Geometry height="{node.height}" width="{node.width}" x="{node.x}" y="{node.y}"/>
          <y:Fill color="{style['fill']}" transparent="false"/>
          <y:BorderStyle color="{style['border']}" type="line" width="1.5"/>
          <y:NodeLabel alignment="center" autoSizePolicy="content" fontFamily="Arial" fontSize="12" fontStyle="plain" hasBackgroundColor="false" hasLineColor="false" horizontalTextPosition="center" iconTextGap="4" modelName="internal" modelPosition="c" textColor="#111827" verticalTextPosition="center" visible="true">{escape(node.label)}</y:NodeLabel>
          <y:Shape type="{style['shape']}"/>
        </y:ShapeNode>
      </data>
    </node>'''


def graphml_edge(edge: Edge) -> str:
    return f'''    <edge id="{escape(edge.edge_id)}" source="{escape(edge.source)}" target="{escape(edge.target)}">
      <data key="d1">
        <y:PolyLineEdge>
          <y:Path sx="0.0" sy="0.0" tx="0.0" ty="0.0"/>
          <y:LineStyle color="#4B5563" type="line" width="1.2"/>
          <y:Arrows source="none" target="standard"/>
          <y:EdgeLabel alignment="center" backgroundColor="#FFFFFF" configuration="AutoFlippingLabel" distance="2.0" fontFamily="Arial" fontSize="11" fontStyle="plain" hasLineColor="false" modelName="centered" modelPosition="center" preferredPlacement="anywhere" ratio="0.5" textColor="#111827" visible="true">{escape(edge.label)}</y:EdgeLabel>
          <y:BendStyle smoothed="false"/>
        </y:PolyLineEdge>
      </data>
    </edge>'''


def generate_graphml(diagram: dict) -> str:
    nodes = '\n'.join(graphml_node(node) for node in diagram['nodes'])
    edges = '\n'.join(graphml_edge(edge) for edge in diagram['edges'])
    return f'''<?xml version="1.0" encoding="UTF-8" standalone="no"?>
<graphml xmlns="http://graphml.graphdrawing.org/xmlns"
    xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"
    xmlns:y="http://www.yworks.com/xml/graphml"
    xmlns:yed="http://www.yworks.com/xml/yed/3"
    xsi:schemaLocation="http://graphml.graphdrawing.org/xmlns http://www.yworks.com/xml/schema/graphml/1.1/ygraphml.xsd">
  <key for="node" id="d0" yfiles.type="nodegraphics"/>
  <key for="edge" id="d1" yfiles.type="edgegraphics"/>
  <graph edgedefault="directed" id="G">
{nodes}
{edges}
  </graph>
</graphml>
'''


def drawio_style(kind: str) -> str:
    style = style_for_kind(kind)
    if style['shape'] == 'ellipse':
        return f'ellipse;whiteSpace=wrap;html=1;fillColor={style["fill"]};strokeColor={style["border"]};fontFamily=Arial;fontSize=12;'
    if style['shape'] == 'diamond':
        return f'rhombus;whiteSpace=wrap;html=1;fillColor={style["fill"]};strokeColor={style["border"]};fontFamily=Arial;fontSize=12;'
    if style['shape'] == 'roundrectangle':
        return f'rounded=1;whiteSpace=wrap;html=1;fillColor={style["fill"]};strokeColor={style["border"]};fontFamily=Arial;fontSize=12;arcSize=12;'
    return f'whiteSpace=wrap;html=1;fillColor={style["fill"]};strokeColor={style["border"]};fontFamily=Arial;fontSize=12;'


def mx_node(node: Node, idx: int) -> str:
    return f'''    <mxCell id="n{idx}" value="{escape(node.label)}" style="{drawio_style(node.kind)}" vertex="1" parent="1">
      <mxGeometry x="{node.x}" y="{node.y}" width="{node.width}" height="{node.height}" as="geometry" />
    </mxCell>'''


def mx_edge(edge: Edge, edge_idx: int, node_map: dict[str, str]) -> str:
    return f'''    <mxCell id="e{edge_idx}" value="{escape(edge.label)}" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;endArrow=block;endFill=1;strokeColor=#4B5563;fontFamily=Arial;fontSize=11;labelBackgroundColor=#FFFFFF;" edge="1" parent="1" source="{node_map[edge.source]}" target="{node_map[edge.target]}">
      <mxGeometry relative="1" as="geometry" />
    </mxCell>'''


def generate_drawio(diagram: dict) -> str:
    node_map: dict[str, str] = {}
    node_xml: list[str] = []
    edge_xml: list[str] = []
    for idx, node in enumerate(diagram['nodes'], start=2):
        node_map[node.node_id] = f'n{idx}'
        node_xml.append(mx_node(node, idx))
    for idx, edge in enumerate(diagram['edges'], start=2 + len(diagram['nodes'])):
        edge_xml.append(mx_edge(edge, idx, node_map))
    cells = '\n'.join(node_xml + edge_xml)
    return f'''<mxfile host="app.diagrams.net" modified="2026-04-23T00:00:00.000Z" agent="Roo" version="24.7.17" type="device">
  <diagram id="{escape(diagram['title']).replace(' ', '_')}" name="{escape(diagram['title'])}">
    <mxGraphModel dx="1280" dy="720" grid="1" gridSize="10" guides="1" tooltips="1" connect="1" arrows="1" fold="1" page="1" pageScale="1" pageWidth="1654" pageHeight="1169" math="0" shadow="0">
      <root>
        <mxCell id="0" />
        <mxCell id="1" parent="0" />
{cells}
      </root>
    </mxGraphModel>
  </diagram>
</mxfile>
'''


def write_diagrams() -> list[Path]:
    outputs: list[Path] = []
    for filename, diagram in ALL_DIAGRAMS:
        graphml_path = DIAGRAM_DIR / f'{filename}.graphml'
        drawio_path = DIAGRAM_DIR / f'{filename}.drawio.xml'
        graphml_path.write_text(generate_graphml(diagram), encoding='utf-8')
        drawio_path.write_text(generate_drawio(diagram), encoding='utf-8')
        outputs.extend([graphml_path, drawio_path])
    return outputs


def set_default_style(document: Document) -> None:
    style = document.styles['Normal']
    style.font.name = 'Arial'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    style.font.size = Pt(11)


def format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, first_line_indent_cm: float | None = None) -> None:
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5
    if first_line_indent_cm is not None:
        paragraph.paragraph_format.first_line_indent = Cm(first_line_indent_cm)
    for run in paragraph.runs:
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        run.font.size = Pt(11)
        run.bold = bold or run.bold


def add_heading_text(document: Document, text: str, level: int = 1) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(12 if level <= 2 else 11)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_para(document: Document, text: str, align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    p = document.add_paragraph(text)
    format_paragraph(p, align=align, first_line_indent_cm=0.9)


def add_body(document: Document, text: str) -> None:
    p = document.add_paragraph(text)
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent_cm=0.9)


def add_item(document: Document, letter: str, title: str, body: str) -> None:
    p = document.add_paragraph()
    p.add_run(f'{letter}. ').bold = True
    p.add_run(title).bold = True
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_body(document, body)


def add_bullets(document: Document, items: Iterable[str]) -> None:
    for item in items:
        p = document.add_paragraph(style='List Bullet')
        p.add_run(item)
        format_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT)


def screenshot_path(name: str) -> Path:
    return SCREENSHOT_DIR / name


def add_screenshot_if_exists(document: Document, filename: str, caption: str, number: int) -> None:
    path = screenshot_path(filename)
    if path.exists():
        document.add_picture(str(path), width=Cm(16.5))
        p = document.add_paragraph(f'Gambar 4.{number} {caption}')
        format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        add_para(document, f'Gambar 4.{number} belum dapat disisipkan karena file {filename} belum tersedia pada folder screenshots. Jalankan proses capture screenshot untuk melengkapinya.')


def create_bab_41() -> Document:
    doc = Document()
    set_default_style(doc)
    add_heading_text(doc, '4.1 Analisis Sistem', 1)
    add_para(doc, f"Sistem yang diaudit merupakan {SYSTEM_PROFILE['domain']} pada {SYSTEM_PROFILE['organization']}. Berdasarkan audit terhadap route, controller, model, migration, helper, filter, dan view, sistem ini dirancang untuk mendukung pencatatan penjualan beras kemasan 5 kg, 10 kg, dan 25 kg secara internal. Proyek tidak menyediakan portal pelanggan publik, sehingga seluruh transaksi dicatat oleh pengguna internal dengan peran admin atau pegawai.")
    add_heading_text(doc, '4.1.1 Analisis Kebutuhan Fungsional', 2)
    add_para(doc, 'Kebutuhan fungsional disusun langsung dari implementasi pada route dan controller yang aktif. Tidak ada role lain selain admin dan pegawai.')
    add_bullets(doc, [
        'Admin dapat login, logout, melihat dashboard admin, dan mengubah profil akun.',
        'Admin dapat mengelola data pengguna melalui modul CRUD pengguna.',
        'Admin dapat mengelola data produk beras dan harga aktif beserta histori harga.',
        'Admin dapat mengelola template transaksi cepat serta pengaturan mode pembatasan penjualan.',
        'Admin dapat melakukan transaksi penjualan, melihat riwayat transaksi, membuka laporan penjualan, dan melihat grafik penjualan.',
        'Pegawai dapat login, logout, melihat dashboard pegawai, mengubah profil, melakukan transaksi manual, melakukan transaksi berbasis template cepat, dan melihat riwayat transaksi yang difilter pada transaksi hari berjalan miliknya.',
    ])
    add_heading_text(doc, '4.1.2 Analisis Kebutuhan Nonfungsional', 2)
    add_bullets(doc, [
        'Keamanan dasar menggunakan session authentication, CSRF filter, validasi input server-side, dan pembatasan akses role berbasis filter.',
        'Konsistensi histori transaksi dijaga dengan snapshot harga pada tabel detail transaksi.',
        'Antarmuka dibangun dengan Bootstrap 5 agar konsisten dan mudah dipakai pada browser modern.',
        'Grafik penjualan disajikan dengan Chart.js untuk membantu pemantauan operasional.',
        'Aplikasi berjalan pada platform web sehingga dapat diakses melalui browser dalam jaringan lokal atau host pengembangan.',
    ])
    add_heading_text(doc, '4.1.3 Kebutuhan Perangkat Keras dan Perangkat Lunak', 2)
    add_bullets(doc, [
        'Perangkat lunak pengembangan: PHP 8.2, Composer, MySQL/MariaDB, dan web browser.',
        'Framework utama: CodeIgniter 4.7 pada sisi backend dan Bootstrap 5 pada sisi antarmuka.',
        'Database yang digunakan adalah MySQL/MariaDB dengan delapan tabel utama yang saling berelasi.',
        'Perangkat keras minimal mencakup komputer atau laptop yang mampu menjalankan web server lokal dan database server secara bersamaan.',
    ])
    add_heading_text(doc, '4.1.4 Analisis Proses Bisnis', 2)
    add_para(doc, 'Alur bisnis aktual dimulai ketika admin atau pegawai melakukan login. Setelah autentikasi berhasil, sistem menampilkan dashboard sesuai role. Admin mempunyai hak untuk menyiapkan data pendukung berupa pengguna, produk, harga, template transaksi, dan pengaturan limit. Saat transaksi dibuat, sistem mengambil harga aktif produk, menghitung total qty, total kilogram, subtotal per kemasan, serta grand total. Jika mode limit aktif dan total kilogram melebihi batas, transaksi ditolak. Jika valid, sistem menyimpan header transaksi, detail item transaksi, dan nomor invoice otomatis. Data transaksi selanjutnya dipakai untuk dashboard, laporan penjualan, dan grafik penjualan.')
    return doc


def create_bab_42() -> Document:
    doc = Document()
    set_default_style(doc)
    add_heading_text(doc, '4.2 Desain Sistem', 1)
    add_para(doc, 'Desain sistem dibuat berdasarkan source code aktual. Struktur DFD dibangun agar konsisten dengan route, controller, service, model, dan tabel database yang benar-benar digunakan oleh aplikasi.')
    add_heading_text(doc, '4.2.1 Diagram Konteks', 2)
    add_para(doc, 'Diagram konteks memodelkan satu proses besar sistem informasi penjualan beras yang berinteraksi dengan dua entitas luar, yaitu admin dan pegawai. Admin mengirimkan data login, data profil, data master, data konfigurasi, dan permintaan laporan; sedangkan pegawai mengirimkan data login, data profil, data transaksi, serta permintaan dashboard.')
    add_para(doc, 'File diagram: diagram/Diagram_Konteks.graphml dan diagram/Diagram_Konteks.drawio.xml.')
    add_heading_text(doc, '4.2.2 DFD Level 0', 2)
    add_para(doc, 'DFD Level 0 memecah sistem menjadi enam proses utama: autentikasi dan profil, kelola pengguna, kelola produk dan harga, kelola template dan mode limit, proses transaksi penjualan, serta dashboard/laporan/grafik. Enam proses ini dipilih karena paling sesuai dengan pemisahan tanggung jawab pada controller dan model aplikasi.')
    add_para(doc, 'File diagram: diagram/DFD_Level0.graphml dan diagram/DFD_Level0.drawio.xml.')
    add_heading_text(doc, '4.2.3 DFD Level 1 Setiap Proses', 2)
    for idx, title in enumerate([
        'P1 Autentikasi dan Profil',
        'P2 Kelola Pengguna',
        'P3 Kelola Produk dan Harga',
        'P4 Kelola Template dan Mode Limit',
        'P5 Proses Transaksi Penjualan',
        'P6 Dashboard, Laporan, dan Grafik',
    ], start=1):
        add_para(doc, f'{title} diuraikan ke dalam DFD Level 1 agar parent-child balancing tetap terjaga. File diagram tersimpan sebagai diagram/DFD_Level1_P{idx}.graphml dan diagram/DFD_Level1_P{idx}.drawio.xml.')
    add_heading_text(doc, '4.2.4 Standarisasi Arus Data', 2)
    add_bullets(doc, [
        'Arus data yang masuk ke proses menggunakan awalan data_.',
        'Arus data yang keluar dari proses menggunakan awalan info_.',
        'Seluruh nama arus data menggunakan snake_case agar konsisten dengan istilah teknis pada dokumen.',
        'Arus antar proses dipakai untuk menjaga konsistensi tanpa memaksa seluruh proses membaca data store yang sama.',
    ])
    add_heading_text(doc, '4.2.5 Audit Balancing Data Store', 2)
    add_table(doc, ['Database', 'Input (data_)', 'Output (info_)', 'Balance'], [list(item) for item in BALANCING])
    add_para(doc, 'Berdasarkan audit final, seluruh data store pada DFD Level 0 memiliki jumlah arus masuk dan arus keluar yang seimbang secara count 1:1 sesuai aturan balancing yang diminta.')
    return doc


def create_bab_43() -> Document:
    doc = Document()
    set_default_style(doc)
    add_heading_text(doc, '4.3 Perancangan Basis Data', 1)
    add_para(doc, 'Perancangan basis data diturunkan dari migration aktif, model, service transaksi, dan query pada controller. Delapan tabel utama dipakai untuk menyimpan akun, produk, harga, template cepat, konfigurasi limit, transaksi, dan detail transaksi.')
    add_heading_text(doc, '4.3.1 Deskripsi Basis Data', 2)
    add_bullets(doc, [
        'Tabel users dipakai untuk menyimpan akun admin dan pegawai.',
        'Tabel products dan product_prices dipakai untuk master produk serta histori harga aktif.',
        'Tabel quick_templates dan quick_template_items dipakai untuk mendukung transaksi cepat.',
        'Tabel sale_limit_settings dipakai untuk pengaturan batas maksimum kilogram per transaksi.',
        'Tabel sales_transactions dan sales_transaction_items dipakai untuk penyimpanan transaksi penjualan beserta snapshot harga.',
    ])
    add_heading_text(doc, '4.3.2 Konseptual Basis Data', 2)
    add_para(doc, 'Secara konseptual, basis data terdiri atas kelompok master, kelompok konfigurasi, dan kelompok transaksi. Kelompok master menyuplai referensi akun dan produk, kelompok konfigurasi mengatur template serta limit, sedangkan kelompok transaksi mencatat penjualan harian.')
    add_para(doc, 'File diagram: diagram/Konseptual_Basis_Data.graphml dan diagram/Konseptual_Basis_Data.drawio.xml.')
    add_heading_text(doc, '4.3.3 Struktur Tabel', 2)
    add_table(doc, ['Tabel', 'Deskripsi', 'Field penting', 'Sumber audit'], [[t['name'], t['description'], ', '.join(t['fields'][:6]) + (' ...' if len(t['fields']) > 6 else ''), t['source']] for t in TABLES])
    add_heading_text(doc, '4.3.4 Relasi Antar Entitas', 2)
    add_table(doc, ['Entitas 1', 'Entitas 2', 'Kardinalitas', 'Foreign key logis'], [list(item) for item in RELATIONSHIPS])
    add_para(doc, 'File diagram ERD tersedia pada diagram/ERD_Chen.graphml, diagram/ERD_Chen.drawio.xml, diagram/ERD_CrowsFoot.graphml, diagram/ERD_CrowsFoot.drawio.xml, diagram/Struktur_Tabel_Relasi.graphml, dan diagram/Struktur_Tabel_Relasi.drawio.xml.')
    return doc


def create_bab_445() -> Document:
    doc = Document()
    set_default_style(doc)
    add_heading_text(doc, 'BAB IV', 1)
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_heading_text(doc, 'HASIL DAN PEMBAHASAN', 1)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading_text(doc, '4.4 Perancangan Antarmuka', 1)
    add_body(doc, 'Perancangan antarmuka pada sistem penjualan beras ini disusun berdasarkan halaman yang benar-benar tersedia pada route, controller, dan view. Uraian difokuskan pada tujuan halaman, komponen utama, susunan tata letak, fitur utama, dan alur penggunaan. Dengan demikian, perancangan antarmuka pada bab ini tidak menambahkan halaman fiktif, melainkan menurunkan seluruh pembahasan dari implementasi aktual sistem.')

    add_heading_text(doc, '4.4.1 Halaman Autentikasi', 2)
    add_item(
        doc,
        'a',
        'Halaman Login Sistem',
        'Halaman login sistem dirancang sebagai pintu masuk autentikasi bagi admin dan pegawai sebelum mengakses modul operasional. Komponen utama pada halaman ini meliputi judul sistem, identitas studi kasus UD Tulus Sari Merta, field username, field password, tombol masuk, serta informasi akun demo. Tata letak halaman dibuat terpusat dengan satu panel login yang ringkas agar pengguna langsung fokus pada proses autentikasi. Fitur utama yang difasilitasi adalah validasi kredensial pengguna dan pembentukan session login. Alur penggunaan dimulai ketika pengguna membuka alamat aplikasi, mengisi username dan password, lalu sistem memvalidasi akun dan mengarahkan pengguna ke dashboard sesuai role.'
    )

    add_heading_text(doc, '4.4.2 Halaman Admin', 2)
    admin_pages = [
        ('Halaman Dashboard Admin', 'Halaman Dashboard Admin dirancang sebagai pusat ringkasan operasional sistem penjualan beras. Komponen utama pada rancangan ini mencakup kartu total transaksi, penjualan hari ini, penjualan bulan ini, status mode limit, jumlah pengguna aktif, jumlah produk aktif, jumlah template aktif, grafik penjualan harian dan bulanan, daftar transaksi terbaru, serta daftar template cepat aktif. Tata letak antarmuka disusun dalam blok statistik di bagian atas, blok grafik di bagian tengah, dan blok tabel maupun kartu informasi pada bagian bawah. Fitur utama yang difasilitasi meliputi pemantauan performa penjualan dan akses cepat ke modul transaksi maupun administrasi. Alur penggunaan dimulai ketika admin berhasil login, membaca indikator utama, lalu memilih modul kerja yang dibutuhkan.'),
        ('Halaman Kelola Pengguna', 'Halaman Kelola Pengguna dirancang untuk mengelola akun admin dan pegawai. Komponen utamanya meliputi tabel daftar pengguna, kolom nama lengkap, username, role, status aktif, serta tombol tambah, edit, dan hapus. Tata letak antarmuka dibuat berbasis tabel agar admin dapat memantau seluruh akun secara ringkas. Fitur utama pada halaman ini adalah CRUD pengguna, validasi username unik, dan perlindungan agar sistem tidak kehilangan admin aktif terakhir. Alur penggunaan dimulai ketika admin membuka daftar pengguna, memilih aksi pada data tertentu, lalu menyimpan perubahan melalui formulir yang tersedia.'),
        ('Halaman Kelola Produk Beras', 'Halaman Kelola Produk Beras dirancang untuk mengelola master produk kemasan 5 kg, 10 kg, dan 25 kg yang dipakai pada transaksi. Komponen utama mencakup tabel daftar produk, kode produk, nama produk, berat kilogram, status aktif, serta tombol tambah, edit, dan hapus. Tata letak halaman menggunakan tabel data sebagai komponen utama dengan formulir terpisah untuk input dan pembaruan data. Fitur utama yang difasilitasi meliputi penambahan produk, perubahan identitas produk, dan pengaturan status produk aktif. Alur penggunaan dimulai dari peninjauan daftar produk, lalu admin melakukan pemeliharaan data sesuai kebutuhan.'),
        ('Halaman Kelola Harga Produk', 'Halaman Kelola Harga Produk dirancang untuk memperbarui harga aktif setiap produk sekaligus menampilkan histori harga. Komponen utama pada halaman ini terdiri atas daftar produk dengan harga saat ini, form input harga baru, effective date, dan tabel histori perubahan harga. Tata letak dibuat dua bagian, yaitu area pembaruan harga aktif dan area histori harga. Fitur utamanya adalah penggantian harga aktif dengan mekanisme histori sehingga transaksi lama tidak berubah. Alur penggunaan dimulai ketika admin memilih produk, memasukkan harga baru dan tanggal efektif, kemudian sistem menyimpan perubahan sebagai histori harga baru.'),
        ('Halaman Kelola Template Transaksi Cepat', 'Halaman Kelola Template Transaksi Cepat dirancang untuk mempercepat pencatatan transaksi berulang. Komponen utama pada halaman ini meliputi daftar template, kode template, nama template, qty default untuk kemasan 5 kg, 10 kg, dan 25 kg, status aktif, serta tombol tambah, edit, dan hapus. Tata letak dirancang berbasis tabel dan form agar proses administrasi template cepat tetap sederhana. Fitur utama yang difasilitasi meliputi pembuatan template baru, pembaruan komposisi qty, dan penonaktifan template yang tidak dipakai. Alur penggunaan dimulai ketika admin mengisi identitas template, menentukan qty tiap kemasan, lalu menyimpan template untuk dipakai pada transaksi berikutnya.'),
        ('Halaman Pengaturan Mode Limit Penjualan', 'Halaman Pengaturan Mode Limit Penjualan dirancang untuk mengaktifkan atau menonaktifkan batas maksimum total kilogram per transaksi. Komponen utama meliputi status saat ini, switch aktivasi mode limit, field batas maksimum kilogram, dan tombol simpan pengaturan. Tata letak dibuat sederhana dalam satu panel konfigurasi agar fokus pengguna tetap pada kontrol aturan transaksi. Fitur utamanya adalah pembatasan jumlah kilogram yang boleh disimpan pada satu transaksi ketika mode limit aktif. Alur penggunaan dimulai saat admin meninjau status limit, mengubah nilai batas maksimum, lalu menyimpan konfigurasi baru.'),
        ('Halaman Laporan Penjualan', 'Halaman Laporan Penjualan dirancang untuk menampilkan rekap transaksi berdasarkan filter tanggal, produk, dan pencatat transaksi. Komponen utama meliputi form filter, kartu ringkasan total transaksi, total kilogram, total pendapatan, tabel daftar transaksi, serta grafik penjualan. Tata letak dibagi menjadi panel filter, panel statistik, panel grafik, dan tabel laporan. Fitur utama yang difasilitasi adalah penyaringan data transaksi dan pembacaan performa penjualan secara periodik. Alur penggunaan dimulai ketika admin menentukan filter laporan, lalu sistem menampilkan data transaksi yang sesuai.'),
        ('Halaman Grafik Penjualan', 'Halaman Grafik Penjualan dirancang untuk menyajikan visualisasi penjualan harian, bulanan, dan tren dua belas bulan. Komponen utama pada halaman ini adalah beberapa kanvas grafik yang dibangun menggunakan Chart.js. Tata letak antarmuka dibuat berbasis kartu agar masing-masing grafik mudah dibaca. Fitur utama yang difasilitasi adalah pemantauan tren penjualan dari sudut pandang visual. Alur penggunaan dimulai ketika admin membuka halaman grafik dan membaca perubahan penjualan dari data yang telah tersimpan di database.'),
        ('Halaman Profil Admin', 'Halaman Profil Admin dirancang untuk mengelola identitas akun pribadi, foto profil, username, dan password. Komponen utama pada halaman ini mencakup data nama lengkap, username, unggah foto profil, password saat ini, password baru, dan konfirmasi password baru. Tata letak dibuat dalam satu halaman profil terpadu agar perubahan data akun dapat dilakukan tanpa berpindah menu. Fitur utamanya adalah pembaruan identitas dan keamanan akun. Alur penggunaan dimulai ketika admin membuka menu profil, memperbarui data yang diperlukan, lalu menyimpan perubahan.'),
    ]
    for index, (title, body) in enumerate(admin_pages):
        add_item(doc, chr(ord('a') + index), title, body)

    add_heading_text(doc, '4.4.3 Halaman Pegawai', 2)
    employee_pages = [
        ('Halaman Dashboard Pegawai', 'Halaman Dashboard Pegawai dirancang sebagai pusat akses cepat ke modul transaksi. Komponen utamanya berupa kartu shortcut menuju transaksi manual, template cepat, dan riwayat transaksi. Tata letak halaman berbentuk kartu navigasi agar pegawai dapat segera memilih pekerjaan operasional. Fitur utama yang difasilitasi adalah navigasi cepat ke transaksi. Alur penggunaan dimulai setelah pegawai login lalu memilih salah satu shortcut yang tersedia.'),
        ('Halaman Transaksi Penjualan Manual', 'Halaman Transaksi Penjualan Manual dirancang untuk mencatat transaksi berdasarkan qty kemasan 5 kg, 10 kg, dan 25 kg. Komponen utama meliputi tanggal transaksi otomatis, nama pelanggan opsional, tabel input qty per kemasan, ringkasan total qty, total kilogram, grand total, status limit pembelian, dan modal konfirmasi transaksi. Tata letak dibuat dua kolom dengan form transaksi pada sisi utama dan panel informasi harga serta limit pada sisi samping. Fitur utama yang difasilitasi adalah perhitungan subtotal otomatis dan validasi sebelum transaksi disimpan. Alur penggunaan dimulai ketika pegawai mengisi qty, memeriksa ringkasan, lalu melakukan konfirmasi penyimpanan transaksi.'),
        ('Halaman Transaksi Berbasis Template Cepat', 'Halaman Transaksi Berbasis Template Cepat dirancang untuk menggunakan template aktif yang telah disusun admin. Komponen utamanya meliputi daftar kartu template aktif, tabel qty kemasan yang terisi otomatis, ringkasan total transaksi, dan modal konfirmasi. Tata letak halaman masih konsisten dengan halaman transaksi manual, namun ditambah panel pemilihan template. Fitur utamanya adalah pengisian qty otomatis dari template untuk mempercepat pencatatan transaksi. Alur penggunaan dimulai ketika pegawai memilih template, memeriksa hasil pengisian qty, lalu melanjutkan konfirmasi transaksi.'),
        ('Halaman Riwayat Transaksi', 'Halaman Riwayat Transaksi dirancang untuk menampilkan daftar transaksi yang telah dicatat. Komponen utama pada halaman ini adalah tabel transaksi, nomor invoice, tanggal, total kilogram, grand total, dan detail item. Tata letak berbasis tabel dipakai agar pegawai mudah membaca daftar transaksi. Fitur utama yang difasilitasi adalah peninjauan histori transaksi dan detail item penjualan. Alur penggunaan dimulai ketika pegawai membuka menu riwayat, lalu memilih transaksi yang ingin dilihat rinciannya.'),
        ('Halaman Profil Pegawai', 'Halaman Profil Pegawai memiliki struktur serupa dengan profil admin karena keduanya memakai modul profil yang sama. Komponen utama mencakup data identitas, username, foto profil, dan perubahan password. Tata letak terpadu membantu pegawai menjaga data akun tetap mutakhir. Fitur utama yang difasilitasi adalah pembaruan data pribadi dan keamanan akun. Alur penggunaan dimulai ketika pegawai membuka menu profil lalu menyimpan perubahan yang diperlukan.'),
    ]
    for index, (title, body) in enumerate(employee_pages):
        add_item(doc, chr(ord('a') + index), title, body)

    add_heading_text(doc, '4.5 Implementasi Sistem', 1)
    add_body(doc, 'Implementasi sistem pada subbab ini menjelaskan halaman yang benar-benar berjalan pada aplikasi penjualan beras. Setiap gambar menggunakan screenshot implementasi nyata yang diambil dari sistem yang berjalan. Apabila screenshot suatu halaman belum tersedia, maka dokumen tetap menuliskan status kekosongan tersebut agar konsisten dengan kondisi aktual aplikasi saat audit dilakukan.')

    add_heading_text(doc, '4.5.1 Implementasi Halaman Sistem', 2)
    implementation_titles = [
        'Halaman Login Sistem',
        'Halaman Dashboard Admin',
        'Halaman Kelola Pengguna',
        'Halaman Kelola Produk Beras',
        'Halaman Kelola Harga Produk',
        'Halaman Kelola Template Transaksi Cepat',
        'Halaman Pengaturan Mode Limit Penjualan',
        'Halaman Laporan Penjualan',
        'Halaman Grafik Penjualan',
        'Halaman Profil Admin',
        'Halaman Dashboard Pegawai',
        'Halaman Transaksi Penjualan Manual',
        'Halaman Transaksi Berbasis Template Cepat',
        'Halaman Riwayat Transaksi',
    ]
    for idx, ((filename, caption), title) in enumerate(zip(SCREENSHOT_PLAN, implementation_titles), start=1):
        add_item(doc, chr(ord('a') + idx - 1), title, f'{title} telah diimplementasikan pada sistem penjualan beras sesuai route, controller, dan view yang diaudit. Komponen yang tampil pada halaman aktual mengikuti fungsi modul masing-masing, sehingga pengguna dapat menjalankan proses bisnis yang tersedia pada sistem secara langsung. Uraian implementasi ini disusun berdasarkan halaman nyata yang dibuka dari aplikasi berjalan.')
        add_screenshot_if_exists(doc, filename, caption.rstrip('.'), idx)
    return doc


def save_document(doc: Document, path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def create_combined_doc() -> Document:
    docs = [create_bab_41(), create_bab_42(), create_bab_43(), create_bab_445()]
    combined = Document()
    set_default_style(combined)
    first = True
    for item in docs:
        if not first:
            combined.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        first = False
        for element in item.element.body:
            combined.element.body.append(element)
    return combined


def write_audit_program() -> Path:
    content = f'''# AUDIT PROGRAM

## Domain Sistem
- {SYSTEM_PROFILE['domain']}
- Studi kasus: {SYSTEM_PROFILE['organization']}

## Role Nyata
- admin
- pegawai

## Modul Nyata
### Halaman Publik
- login

### Modul Auth dan Profil
- login/logout berbasis session
- ubah profil, ubah username, upload foto profil, ubah password

### Modul Admin
- dashboard admin
- kelola pengguna
- kelola produk
- kelola harga
- kelola template transaksi cepat
- mode pembatasan penjualan
- laporan penjualan
- grafik penjualan

### Modul Pegawai
- dashboard pegawai
- transaksi manual
- transaksi template cepat
- riwayat transaksi hari berjalan
- profil

## Fitur Nyata
{chr(10).join('- ' + item for item in FEATURE_AUDIT['master_modules'] + FEATURE_AUDIT['transaction_modules'] + FEATURE_AUDIT['reporting_modules'] + FEATURE_AUDIT['profile_modules'])}

## Tabel Nyata
{chr(10).join(f"- {t['name']}: {t['description']}" for t in TABLES)}

## Relasi Utama
{chr(10).join(f"- {a} -> {b} ({c}, fk: {d})" for a, b, c, d in RELATIONSHIPS)}

## Alur Bisnis Utama
1. Pengguna login ke sistem.
2. Sistem memvalidasi akun aktif berdasarkan tabel users.
3. Admin menyiapkan data master produk, harga, template, pengguna, dan limit penjualan.
4. Admin atau pegawai melakukan input transaksi manual atau dari template.
5. Sistem mengambil harga aktif, menghitung subtotal, total kg, grand total, dan memvalidasi limit penjualan.
6. Sistem menyimpan header transaksi dan detail transaksi dengan snapshot harga.
7. Dashboard, laporan, dan grafik mengambil data transaksi yang telah tersimpan.

## Catatan Integrasi atau Otomatisasi
- Nomor invoice digenerate otomatis oleh service transaksi.
- Grafik dibentuk dengan Chart.js dari agregasi penjualan harian dan bulanan.
- Folder `/contoh` berisi acuan format BAB 4, terutama pada dokumen desain DFD, perancangan basis data, dan pola penulisan 4.4 serta 4.5.
'''
    path = AUDIT_DIR / 'AUDIT_PROGRAM.md'
    path.write_text(content, encoding='utf-8')
    return path


def write_audit_balancing() -> Path:
    lines = [
        '# AUDIT DFD BALANCING',
        '',
        '## Entitas Luar',
        '- Admin',
        '- Pegawai',
        '',
        '## Proses Level 0',
        '- P1 Autentikasi dan Profil',
        '- P2 Kelola Pengguna',
        '- P3 Kelola Produk dan Harga',
        '- P4 Kelola Template dan Mode Limit',
        '- P5 Proses Transaksi Penjualan',
        '- P6 Dashboard, Laporan, dan Grafik',
        '',
        '## Mapping Parent-Child Level 1',
        '- P1 -> Level 1 P1: validasi login dan pembaruan profil',
        '- P2 -> Level 1 P2: simpan dan tampilkan data pengguna',
        '- P3 -> Level 1 P3: simpan/tampilkan produk dan harga',
        '- P4 -> Level 1 P4: simpan/tampilkan template serta limit penjualan',
        '- P5 -> Level 1 P5: terima, validasi, dan simpan transaksi',
        '- P6 -> Level 1 P6: susun laporan, grafik, dan dashboard',
        '',
        '## Rename Arus Data Lama ke Standar Final',
        '- informasi produk -> info_produk',
        '- informasi harga -> info_harga_produk',
        '- data limit -> data_limit_penjualan',
        '- informasi laporan -> info_laporan_penjualan',
        '- informasi dashboard -> info_dashboard_admin / info_dashboard_pegawai',
        '',
        '## Tabel Balancing Data Store',
        '',
        '| Database | Input (data_) | Output (info_) | Balance |',
        '|---|---|---|---|',
    ]
    for store, data_in, data_out, balance in BALANCING:
        lines.append(f'| {store} | {data_in} | {data_out} | {balance} |')
    lines.extend([
        '',
        '## Konsistensi Antar Diagram',
        '- Diagram konteks selaras dengan DFD Level 0 karena seluruh alur utama admin dan pegawai dipertahankan.',
        '- DFD Level 0 selaras dengan seluruh DFD Level 1 karena setiap parent process memiliki child diagram tersendiri.',
        '- Data store pada Level 0 dijaga seimbang secara count 1:1.',
        '',
        '## Hal yang Direvisi',
        '- Akses laporan dan grafik ditempatkan pada proses P6 agar tetap konsisten dengan controller Reports dan Charts.',
        '- Akses data pendukung transaksi dialirkan sebagai arus antar proses dari P3 dan P4 ke P5 untuk menjaga disiplin balancing store.',
        '- Proses dashboard dipisahkan dari autentikasi agar tidak mencampur fungsi login dengan pelaporan operasional.',
    ])
    path = AUDIT_DIR / 'AUDIT_DFD_BALANCING.md'
    path.write_text('\n'.join(lines), encoding='utf-8')
    return path


def write_audit_doc() -> Path:
    content = '''# AUDIT DOC BAB4

## Sumber Isi Tiap Subbab
- BAB 4.1 disusun dari README, route, controller, filter, helper, migration, model, dan view.
- BAB 4.2 disusun dari hasil audit proses bisnis, role, dan tabel yang kemudian dimodelkan menjadi Diagram Konteks, DFD Level 0, dan DFD Level 1.
- BAB 4.3 disusun dari migration utama, model, foreign key, dan service transaksi.
- BAB 4.4 disusun dari halaman view yang benar-benar ada pada folder app/Views.
- BAB 4.5 disusun dari screenshot implementasi nyata yang diambil dari halaman yang berjalan.

## Halaman yang Dipakai dari Program
- auth/login
- dashboard/index
- users/index dan users/form
- products/index dan products/form
- prices/index
- templates/index dan templates/form
- sale_limit/index
- sales/create dan sales/index
- reports/index
- charts/index
- profile/index

## Halaman yang Tidak Dimasukkan karena Tidak Ada di Program
- register publik
- halaman checkout customer
- halaman pembayaran digital
- halaman stok gudang
- halaman produksi/penggilingan
- halaman notifikasi pihak ketiga

## Alasan Penyusunan 4.4 dan 4.5
- Antarmuka dibagi berdasarkan struktur menu admin dan pegawai yang benar-benar tampil pada sidebar.
- Implementasi disusun mengikuti urutan akses yang paling logis: login, dashboard, modul master, modul transaksi, modul laporan, dan profil.
- Struktur penulisan 4.4 dan 4.5 diselaraskan dengan pola contoh pada folder `/contoh`, terutama bentuk subbab, item huruf, dan gaya uraian akademik per halaman.
'''
    path = AUDIT_DIR / 'AUDIT_DOC_BAB4.md'
    path.write_text(content, encoding='utf-8')
    return path


def write_manifest(existing: list[Path]) -> Path:
    lines = ['# OUTPUT MANIFEST', '']
    for path in sorted(existing, key=lambda p: p.as_posix()):
        rel = path.relative_to(ROOT).as_posix()
        desc = 'File output akhir hasil generator'
        if rel.startswith('diagram/'):
            desc = 'File diagram final dalam format graphml atau drawio xml'
        elif rel.startswith('docs/'):
            desc = 'Dokumen BAB 4 final dalam format docx'
        elif rel.startswith('audit/'):
            desc = 'Dokumen audit konsistensi dan manifest'
        elif rel.startswith('screenshots/'):
            desc = 'Screenshot implementasi nyata dari aplikasi berjalan'
        lines.append(f'- `{rel}` : {desc}')
    path = AUDIT_DIR / 'OUTPUT_MANIFEST.md'
    path.write_text('\n'.join(lines) + '\n', encoding='utf-8')
    return path


def ensure_dirs() -> None:
    for directory in [DIAGRAM_DIR, DOCS_DIR, AUDIT_DIR, SCREENSHOT_DIR, OUTPUT_AUDIT_DIR, OUTPUT_DIAGRAM_DIR, OUTPUT_DOCX_DIR]:
        directory.mkdir(parents=True, exist_ok=True)


def main() -> None:
    ensure_dirs()
    outputs: list[Path] = []
    outputs.extend(write_diagrams())

    doc41 = save_document(create_bab_41(), DOCS_DIR / 'BAB_4_1_Analisis.docx')
    doc42 = save_document(create_bab_42(), DOCS_DIR / 'BAB_4_2_Desain.docx')
    doc43 = save_document(create_bab_43(), DOCS_DIR / 'BAB_4_3_Perancangan_Basis_Data.docx')
    doc445 = save_document(create_bab_445(), DOCS_DIR / '4.4.4.5.docx')
    doc_all = save_document(create_combined_doc(), DOCS_DIR / 'BAB_4_1_sampai_4_5.docx')
    outputs.extend([doc41, doc42, doc43, doc445, doc_all])

    outputs.append(write_audit_program())
    outputs.append(write_audit_balancing())
    outputs.append(write_audit_doc())

    for filename, _caption in SCREENSHOT_PLAN:
        path = screenshot_path(filename)
        if path.exists():
            outputs.append(path)

    outputs.append(write_manifest(outputs))
    print('Artifacts generated successfully.')


if __name__ == '__main__':
    main()
